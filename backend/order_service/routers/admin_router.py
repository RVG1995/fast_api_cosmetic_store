"""API эндпоинты сервиса заказов для администрирования заказов.
Предоставляет интерфейсы только для администраторов с расширенными функциями управления."""

import os
import logging
from typing import Optional, List, Dict, Any
from datetime import datetime
import hashlib
from dotenv import load_dotenv
from fastapi import APIRouter, Depends, HTTPException, Query, Path, status, Body
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
import pandas as pd
import io
# Импортируем необходимые библиотеки
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER

# Добавляем импорт для поддержки кириллицы
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from database import get_db
from models import OrderStatusModel, OrderStatusHistoryModel, PromoCodeModel, DeliveryInfoModel,OrderModel
from schemas import (
    OrderUpdate, OrderResponse, OrderDetailResponse, 
    OrderStatusHistoryCreate, PaginatedResponse, OrderStatistics, BatchStatusUpdate,
    OrderItemsUpdate, OrderItemsUpdateResponse, PromoCodeResponse, OrderResponseWithPromo, OrderDetailResponseWithPromo,
    AdminOrderCreate, OrderFilterParams, OrderCreate, DeliveryInfoUpdate
)
from services import (
    create_order, get_order_by_id, get_orders, update_order, 
    change_order_status, get_order_statistics, get_order_statistics_by_date,
    update_order_items
)
from dependencies import (
    get_admin_user, get_order_filter_params,
    check_products_availability
)
from cache import (
    get_cached_order, cache_order, invalidate_order_cache,
    get_cached_order_statistics, cache_order_statistics, invalidate_statistics_cache,
    get_cached_orders_list, cache_orders_list, CacheKeys, get_cached_data, set_cached_data,
    DEFAULT_CACHE_TTL
)
from notification_api import check_notification_settings
from sqlalchemy import select
from sqlalchemy.orm import selectinload

# Загрузка переменных окружения
load_dotenv()

# Настройка логирования
logger = logging.getLogger("admin_order_router")

# Получение URL сервиса продуктов
PRODUCT_SERVICE_URL = os.getenv("PRODUCT_SERVICE_URL", "http://localhost:8001")

# Создание роутера для админ-панели
router = APIRouter(
    prefix="/admin/orders",
    tags=["admin_orders"],
    responses={404: {"description": "Not found"}},
)

# Переопределяем get_order_by_id с подгрузкой связей (если не было)
async def get_order_with_all_relations(session, order_id):
    result = await session.execute(
        select(OrderModel)
        .options(
            selectinload(OrderModel.delivery_info),
            selectinload(OrderModel.items),
            selectinload(OrderModel.status),
            selectinload(OrderModel.status_history),
            selectinload(OrderModel.promo_code),
        )
        .where(OrderModel.id == order_id)
    )
    return result.scalar_one_or_none()

# Вспомогательная функция для безопасного получения данных заказа
async def get_safe_order_response(session, order_id):
    """
    Безопасно получает данные заказа и создаёт ответ без асинхронной загрузки вложенных объектов
    """
    # Загружаем заказ со всеми связанными данными
    loaded_order = await get_order_with_all_relations(session, order_id)
    if not loaded_order:
        return None
        
    # Собираем основные данные о заказе в словарь
    order_dict = {
        "id": loaded_order.id,
        "user_id": loaded_order.user_id,
        "status_id": loaded_order.status_id,
        "status": {
            "id": loaded_order.status.id,
            "name": loaded_order.status.name,
            "description": loaded_order.status.description,
            "color": loaded_order.status.color,
            "allow_cancel": loaded_order.status.allow_cancel,
            "is_final": loaded_order.status.is_final,
            "sort_order": loaded_order.status.sort_order
        },
        "created_at": loaded_order.created_at,
        "updated_at": loaded_order.updated_at,
        "total_price": loaded_order.total_price,
        "promo_code_id": loaded_order.promo_code_id,
        "discount_amount": loaded_order.discount_amount,
        "full_name": loaded_order.full_name,
        "email": loaded_order.email,
        "phone": loaded_order.phone,
        "delivery_address": loaded_order.delivery_address,
        "comment": loaded_order.comment,
        "is_paid": loaded_order.is_paid,
        "items": [
            {
                "id": item.id,
                "order_id": item.order_id,
                "product_id": item.product_id,
                "product_name": item.product_name,
                "product_price": item.product_price,
                "quantity": item.quantity,
                "total_price": item.total_price
            } 
            for item in loaded_order.items
        ],
        "order_number": f"{loaded_order.id}-{loaded_order.created_at.year}",
        "personal_data_agreement": getattr(loaded_order, 'personal_data_agreement', None),
        "is_payment_on_delivery": getattr(loaded_order, 'is_payment_on_delivery', True),
        "receive_notifications": getattr(loaded_order, 'receive_notifications', None)
    }
    
    # Получаем промокод отдельным запросом если необходимо
    if loaded_order.promo_code_id:
        promo_result = await session.execute(select(PromoCodeModel).where(PromoCodeModel.id == loaded_order.promo_code_id))
        promo = promo_result.scalar_one_or_none()
        if promo:
            order_dict["promo_code"] = {
                "id": promo.id,
                "code": promo.code,
                "discount_percent": promo.discount_percent,
                "discount_amount": promo.discount_amount,
                "valid_until": promo.valid_until,
                "is_active": promo.is_active,
                "created_at": promo.created_at,
                "updated_at": promo.updated_at
            }
    
    # Получаем информацию о доставке отдельным запросом
    delivery_query = select(DeliveryInfoModel).where(DeliveryInfoModel.order_id == loaded_order.id)
    delivery_result = await session.execute(delivery_query)
    delivery_info = delivery_result.scalar_one_or_none()
    
    if delivery_info:
        order_dict["delivery_info"] = {
            "id": delivery_info.id,
            "order_id": delivery_info.order_id,
            "delivery_type": delivery_info.delivery_type,
            "boxberry_point_id": delivery_info.boxberry_point_id,
            "boxberry_point_address": delivery_info.boxberry_point_address,
            "delivery_cost": delivery_info.delivery_cost,
            "tracking_number": delivery_info.tracking_number,
            "label_url_boxberry": delivery_info.label_url_boxberry,
            "status_in_delivery_service": delivery_info.status_in_delivery_service
        }
    
    # Получаем историю статусов отдельным запросом
    status_history_query = select(OrderStatusHistoryModel).where(
        OrderStatusHistoryModel.order_id == loaded_order.id
    ).options(selectinload(OrderStatusHistoryModel.status))
    status_history_result = await session.execute(status_history_query)
    status_history_items = status_history_result.scalars().all()
    
    # Формируем историю статусов вручную
    order_dict["status_history"] = []
    for history in status_history_items:
        history_dict = {
            "id": history.id,
            "order_id": history.order_id,
            "status_id": history.status_id,
            "changed_at": history.changed_at,
            "changed_by_user_id": history.changed_by_user_id,
            "notes": history.notes,
        }
        
        if history.status:
            history_dict["status"] = {
                "id": history.status.id,
                "name": history.status.name,
                "description": history.status.description,
                "color": history.status.color,
                "allow_cancel": history.status.allow_cancel,
                "is_final": history.status.is_final,
                "sort_order": history.status.sort_order
            }
        else:
            history_dict["status"] = None
            
        order_dict["status_history"].append(history_dict)
    
    # Создаем и возвращаем ответ из словаря
    return OrderDetailResponseWithPromo.model_validate(order_dict)

@router.get("", response_model=PaginatedResponse,dependencies=[Depends(get_admin_user)])
async def list_all_orders(
    filters = Depends(get_order_filter_params),
    session: AsyncSession = Depends(get_db),
):
    """
    Получение списка всех заказов (только для администраторов).
    
    Поддерживает фильтрацию по различным параметрам.
    """
    try:
        # Создаем ключ для кэша на основе параметров фильтрации
        filter_params = filters.model_dump_json()
        cache_key = hashlib.md5(filter_params.encode()).hexdigest()
        
        # Пытаемся получить данные из кэша
        cached_orders = await get_cached_orders_list(cache_key)
        if cached_orders:
            logger.info("Данные о всех заказах получены из кэша по ключу %s", cache_key)
            return cached_orders
        
        # Если данных нет в кэше, получаем из БД
        orders, total = await get_orders(
            session=session,
            filters=filters
        )
        
        # Формируем ответ с преобразованием моделей в схемы
        order_responses = []
        for order in orders:
            # Создаем базовый объект OrderResponse без промокода
            order_response = OrderResponse.model_validate(order)
            
            # Если у заказа есть промокод, создаем OrderResponseWithPromo
            if order.promo_code_id:
                # Создаем расширенный ответ с информацией о промокоде
                with_promo = OrderResponseWithPromo(**order_response.model_dump())
                
                try:
                    # Получаем промокод и устанавливаем его
                    promo_code = await session.get(PromoCodeModel, order.promo_code_id)
                    if promo_code:
                        with_promo.promo_code = PromoCodeResponse.model_validate(promo_code)
                        logger.info("Для заказа %s загружен промокод %s (в админ списке)", order.id, promo_code.code)
                except (ValueError, AttributeError, KeyError) as e:
                    logger.warning("Не удалось загрузить промокод %s для заказа %s: %s", order.promo_code_id, order.id, str(e))
                
                order_responses.append(with_promo)
            else:
                # Заказ без промокода, используем обычный OrderResponse
                order_responses.append(order_response)
        
        response = PaginatedResponse(
            items=order_responses,
            total=total,
            page=filters.page,
            size=filters.size,
            pages=0  # Будет вычислено в валидаторе
        )
        
        # Кэшируем результат
        await cache_orders_list(cache_key, response)
        logger.info("Данные о всех заказах добавлены в кэш по ключу %s", cache_key)
        
        return response
    except Exception as e:
        logger.error("Ошибка при получении списка заказов: %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Произошла ошибка при получении списка заказов",
        ) from e

@router.get("/statistics", response_model=OrderStatistics, dependencies=[Depends(get_admin_user)])
async def get_admin_orders_statistics(
    session: AsyncSession = Depends(get_db),
):
    """
    Получение статистики по всем заказам (только для администраторов)
    """
    try:
        # Пытаемся получить данные из кэша
        cached_statistics = await get_cached_order_statistics()
        if cached_statistics:
            logger.info("Статистика всех заказов получена из кэша")
            return cached_statistics
        
        # Если данных нет в кэше, получаем из БД
        statistics = await get_order_statistics(session)
        
        # Кэшируем результат
        await cache_order_statistics(statistics)
        logger.info("Статистика всех заказов добавлена в кэш")
        
        return statistics
    except Exception as e:
        logger.error("Ошибка при получении статистики заказов: %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Ошибка при получении статистики заказов"
        ) from e

@router.get("/statistics/report", response_model=OrderStatistics, dependencies=[Depends(get_admin_user)])
async def get_admin_orders_statistics_by_date(
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
    session: AsyncSession = Depends(get_db),
):
    """
    Получение статистики по заказам за указанный период (только для администраторов)
    
    - **date_from**: Дата начала периода в формате YYYY-MM-DD
    - **date_to**: Дата окончания периода в формате YYYY-MM-DD
    """
    try:
        # Формируем ключ кэша на основе параметров запроса
        cache_key = f"{CacheKeys.ORDER_REPORTS_PREFIX}{date_from or 'all'}:{date_to or 'all'}"
        
        # Пытаемся получить данные из кэша
        cached_statistics = await get_cached_data(cache_key)
        if cached_statistics:
            logger.info("Статистика заказов за период получена из кэша")
            return cached_statistics
        
        # Если данных нет в кэше, получаем из БД
        statistics = await get_order_statistics_by_date(session, date_from, date_to)
        
        # Кэшируем результат на короткое время (5 минут)
        await set_cached_data(cache_key, statistics, 300)
        logger.info("Статистика заказов за период добавлена в кэш")
        
        return statistics
    except Exception as e:
        logger.error("Ошибка при получении статистики заказов за период: %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Ошибка при получении статистики заказов за период"
        ) from e

@router.get("/{order_id}", response_model=OrderDetailResponseWithPromo, dependencies=[Depends(get_admin_user)])
async def get_order_admin(
    order_id: int = Path(..., ge=1),
    session: AsyncSession = Depends(get_db),
):
    """
    Получение информации о заказе по ID (только для администраторов)
    
    - **order_id**: ID заказа
    """
    try:
        # Пытаемся получить данные из кэша
        cached_order = await get_cached_order(order_id, admin=True)
        if cached_order:
            logger.info("Данные о заказе %s получены из кэша (админ)", order_id)
            return cached_order
            
        # Получаем заказ
        order = await get_order_by_id(session, order_id)
        
        if not order:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Заказ с ID {order_id} не найден",
            )
        
        # Преобразуем модель в схему без промокода
        order_response = OrderDetailResponse.model_validate(order)
        
        # Создаем расширенный ответ с возможностью добавления промокода
        response_with_promo = OrderDetailResponseWithPromo(**order_response.model_dump())
        
        # Вручную обрабатываем промокод
        if order.promo_code_id:
            try:
                promo_code = await session.get(PromoCodeModel, order.promo_code_id)
                if promo_code:
                    response_with_promo.promo_code = PromoCodeResponse.model_validate(promo_code)
                    logger.info("Для заказа %s загружен промокод %s (админ)", order.id, promo_code.code)
            except (ValueError, AttributeError, KeyError) as e:
                logger.warning("Не удалось загрузить промокод %s для заказа %s: %s", order.promo_code_id, order.id, str(e))
        
        # Кэшируем результат с админским ключом
        cache_key = f"{CacheKeys.ORDER_PREFIX}{order_id}:admin"
        await cache_order(order_id, response_with_promo, admin=True, cache_key=cache_key)
        logger.info("Данные о заказе %s добавлены в кэш (админ)", order_id)
        
        return response_with_promo
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Ошибка при получении информации о заказе (админ): %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Произошла ошибка при получении информации о заказе",
        ) from e 

@router.put("/{order_id}", response_model=OrderDetailResponseWithPromo)
async def update_order_admin(
    order_id: int = Path(..., ge=1),
    order_data: OrderUpdate = Body(...),
    current_user: Dict[str, Any] = Depends(get_admin_user),
    session: AsyncSession = Depends(get_db)
):
    """
    Комплексное обновление заказа и информации о доставке (только для администраторов).
    
    - **order_id**: ID заказа
    - **order_data**: Данные для обновления заказа включая доставку, товары и статус
    """
    try:
        logger.info("Получен запрос на комплексное обновление заказа %s", order_id)
        logger.info("Данные заказа: %s", order_data.model_dump())
        
        # Проверяем существование заказа
        updated_order = await get_order_with_all_relations(session, order_id)
        if not updated_order:
            raise HTTPException(status_code=404, detail="Заказ не найден")
        
        # --- ЗАПРЕТ НА ИЗМЕНЕНИЯ ДЛЯ ОПЛАЧЕННЫХ ---
        if updated_order.is_paid:
            allowed_fields = {'comment', 'status_id'}
            forbidden_changes = [
                f for f in order_data.model_fields_set
                if f not in allowed_fields
            ]
            if forbidden_changes:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Редактирование невозможно: заказ уже оплачен. Разрешено менять только комментарий и статус."
                )
        
        # Проверка статуса заказа - запрещаем редактирование для определенных статусов
        non_editable_statuses = ["Отправлен", "Доставлен", "Отменен", "Оплачен"]
        if updated_order.status and updated_order.status.name in non_editable_statuses:
            if order_data.status_id is None:  # Если не меняем статус, а только редактируем заказ
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST, 
                    detail=f"Редактирование заказа невозможно в статусе '{updated_order.status.name}'"
                )
            else:
                # Если меняем статус, то другие изменения не применяем
                logger.warning("Заказ в статусе %s, будет изменен только статус", updated_order.status.name)
                
                # Проверяем, отличается ли новый статус от текущего
                if order_data.status_id == updated_order.status_id:
                    logger.info("Статус заказа не изменился, пропускаем обновление")
                    # Используем тот же подход что и в конце функции
                    # Вручную запрашиваем заказ со всеми связанными данными и формируем безопасный ответ
                    return await get_safe_order_response(session, updated_order.id)
                
                # Проверяем существование нового статуса
                new_status = await session.get(OrderStatusModel, order_data.status_id)
                if not new_status:
                    raise ValueError(f"Статус с ID {order_data.status_id} не найден")
                
                # Меняем только статус
                updated_order.status_id = order_data.status_id
                updated_order.updated_at = datetime.utcnow()
                
                # Добавляем запись в историю статусов
                status_note = order_data.comment if order_data.comment else "Статус обновлен администратором"
                await OrderStatusHistoryModel.add_status_change(
                    session=session,
                    order_id=updated_order.id,
                    status_id=order_data.status_id,
                    changed_by_user_id=current_user["user_id"],
                    notes=status_note
                )
                
                await session.commit()
                
                # Инвалидируем кэш
                await invalidate_order_cache(updated_order.id)
                await invalidate_statistics_cache()
                
                # Используем безопасную функцию для создания ответа
                return await get_safe_order_response(session, updated_order.id)
        
        # Проверка оплаченных заказов - запрещаем редактирование товаров для оплаченных заказов
        if updated_order.is_paid and (
            getattr(order_data, "items_to_add", None) or
            getattr(order_data, "items_to_update", None) or
            getattr(order_data, "items_to_remove", None)
        ):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Редактирование товаров невозможно для оплаченных заказов"
            )
        
        # Изменения внесены в заказ
        changes_made = False
        
        # Обновляем основные поля заказа
        if hasattr(order_data, 'full_name') and order_data.full_name is not None:
            updated_order.full_name = order_data.full_name
            changes_made = True
        if hasattr(order_data, 'email') and order_data.email is not None:
            updated_order.email = order_data.email
            changes_made = True
        if hasattr(order_data, 'phone') and order_data.phone is not None:
            updated_order.phone = order_data.phone
            changes_made = True
        if hasattr(order_data, 'delivery_address') and order_data.delivery_address is not None:
            updated_order.delivery_address = order_data.delivery_address
            changes_made = True
        if hasattr(order_data, 'comment') and order_data.comment is not None:
            updated_order.comment = order_data.comment
            changes_made = True
        if hasattr(order_data, 'is_paid') and order_data.is_paid is not None:
            updated_order.is_paid = order_data.is_paid
            changes_made = True
        if hasattr(order_data, 'is_payment_on_delivery') and order_data.is_payment_on_delivery is not None:
            updated_order.is_payment_on_delivery = order_data.is_payment_on_delivery
            changes_made = True
        
        # Обновляем информацию о доставке
        if hasattr(order_data, 'delivery_info') and order_data.delivery_info is not None:
            # Получаем информацию о доставке для заказа
            delivery_info_query = select(DeliveryInfoModel).filter(DeliveryInfoModel.order_id == updated_order.id)
            delivery_info_result = await session.execute(delivery_info_query)
            delivery_info = delivery_info_result.scalars().first()
            
            # Если информация о доставке уже существует, обновляем её
            if delivery_info:
                if order_data.delivery_info.delivery_type is not None:
                    delivery_info.delivery_type = order_data.delivery_info.delivery_type
                    changes_made = True
                if order_data.delivery_info.boxberry_point_id is not None:
                    delivery_info.boxberry_point_id = order_data.delivery_info.boxberry_point_id
                    changes_made = True
                if order_data.delivery_info.boxberry_point_address is not None:
                    delivery_info.boxberry_point_address = order_data.delivery_info.boxberry_point_address
                    changes_made = True
                if order_data.delivery_info.delivery_cost is not None:
                    # Если изменяется стоимость доставки
                    old_delivery_cost = delivery_info.delivery_cost
                    delivery_info.delivery_cost = order_data.delivery_info.delivery_cost
                    
                    # Пересчитываем общую сумму заказа с новой стоимостью доставки
                    if old_delivery_cost != order_data.delivery_info.delivery_cost:
                        # Вычитаем старую стоимость доставки и добавляем новую
                        updated_order.total_price = updated_order.total_price - old_delivery_cost + order_data.delivery_info.delivery_cost
                        logger.info(f"Пересчитана общая сумма заказа: старая доставка {old_delivery_cost}, новая доставка {order_data.delivery_info.delivery_cost}, новая сумма {updated_order.total_price}")
                    
                    changes_made = True
                if order_data.delivery_info.tracking_number is not None:
                    delivery_info.tracking_number = order_data.delivery_info.tracking_number
                    changes_made = True
                    logger.info("Обновлен трек-номер для заказа %s: %s", updated_order.id, order_data.delivery_info.tracking_number)
                if order_data.delivery_info.label_url_boxberry is not None:
                    delivery_info.label_url_boxberry = order_data.delivery_info.label_url_boxberry
                    changes_made = True
        
        # Обновляем статус заказа, если он предоставлен
        if order_data.status_id is not None and order_data.status_id != updated_order.status_id:
            # Проверяем существование статуса
            status_order = await session.get(OrderStatusModel, order_data.status_id)
            if not status_order:
                logger.error("Статус с ID %s не найден", order_data.status_id)
                raise ValueError(f"Статус с ID {order_data.status_id} не найден")
            
            # Устанавливаем новый статус
            updated_order.status_id = order_data.status_id
            changes_made = True
            
            # Добавляем запись в историю статусов
            status_note = order_data.comment if order_data.comment else "Статус обновлен администратором"
            await OrderStatusHistoryModel.add_status_change(
                session=session,
                order_id=updated_order.id,
                status_id=order_data.status_id,
                changed_by_user_id=current_user["user_id"],
                notes=status_note
            )
        
        # Обновляем товары в заказе, если они предоставлены
        if (hasattr(order_data, 'items_to_add') and order_data.items_to_add) or \
           (hasattr(order_data, 'items_to_update') and order_data.items_to_update) or \
           (hasattr(order_data, 'items_to_remove') and order_data.items_to_remove):
            
            # Подготавливаем данные для обновления
            items_to_add = [item.model_dump() for item in order_data.items_to_add] if order_data.items_to_add else []
            items_to_update = order_data.items_to_update if order_data.items_to_update else {}
            items_to_remove = order_data.items_to_remove if order_data.items_to_remove else []
            
            # Проверяем, есть ли изменения в товарах
            if items_to_add or items_to_update or items_to_remove:
                # Вызываем сервисную функцию для обновления товаров
                updated_order = await update_order_items(
                    order_id=updated_order.id,
                    items_to_add=items_to_add,
                    items_to_update=items_to_update,
                    items_to_remove=items_to_remove,
                    session=session,
                    user_id=current_user["user_id"]
                )
                
                if not updated_order:
                    raise HTTPException(status_code=400, detail="Ошибка при обновлении элементов заказа")
                
                # Обновляем текущий заказ данными о товарах
                updated_order = await get_order_with_all_relations(session, updated_order.id)
                changes_made = True
        
        if changes_made:
            # Обновляем дату обновления заказа
            updated_order.updated_at = datetime.utcnow()
            
            # Коммитим изменения
            await session.commit()
            
            # Инвалидируем кэш заказа
            await invalidate_order_cache(updated_order.id)
            # Инвалидируем кэш статистики, если изменился статус
            if order_data.status_id is not None and order_data.status_id != updated_order.status_id:
                await invalidate_statistics_cache()
            logger.info("Кэш заказа %s инвалидирован после обновления", updated_order.id)
        else:
            logger.info("Не обнаружено изменений для заказа %s", updated_order.id)
        
        # Используем вспомогательную функцию для безопасного создания ответа
        return await get_safe_order_response(session, updated_order.id)
    except ValueError as e:
        logger.error("Ошибка при обновлении заказа: %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(e),
        ) from e
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Непредвиденная ошибка при обновлении заказа: %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Произошла ошибка при обновлении заказа: {str(e)}",
        ) from e

@router.post("/{order_id}/status", response_model=OrderResponse)
async def update_order_status(
    order_id: int = Path(..., ge=1),
    status_data: dict = Body(...),
    current_user: Dict[str, Any] = Depends(get_admin_user),
    session: AsyncSession = Depends(get_db)
):
    """
    Обновление статуса заказа (только для администраторов).
    Упрощенный эндпоинт для обновления только статуса.
    
    - **order_id**: ID заказа
    - **status_data**: Данные о новом статусе:
      - status_id: ID нового статуса
      - comment: Комментарий к изменению статуса (опционально)
    """
    try:
        logger.info("Запрос на обновление статуса заказа. ID: %s, данные: %s", order_id, status_data)
        
        # Проверяем наличие обязательного поля status_id
        if not status_data.get("status_id"):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Не указан ID статуса",
            )
        
        # Создаем объект OrderUpdate только с нужными полями
        status_id = status_data.get("status_id")
        comment = status_data.get("comment")
        
        # Проверяем, что статус существует
        order_status = await session.get(OrderStatusModel, status_id)
        if not order_status:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Статус с ID {status_id} не найден",
            )
        
        # Получаем текущий заказ
        order = await get_order_by_id(session, order_id)
        if not order:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Заказ с ID {order_id} не найден",
            )
        
        # Обновляем статус заказа
        order.status_id = status_id
        session.add(order)
        
        # Добавляем запись в историю статусов
        status_note = comment if comment else "Статус обновлен администратором"
        await OrderStatusHistoryModel.add_status_change(
            session=session,
            order_id=order_id,
            status_id=status_id,
            changed_by_user_id=current_user["user_id"],
            notes=status_note
        )
        
        # Получаем информацию о новом статусе
        new_status = await session.get(OrderStatusModel, status_id)
        new_status_name = new_status.name if new_status else "Неизвестный"
        
        # Если новый статус - "Оплачен", автоматически устанавливаем флаг is_paid=True
        if new_status and new_status.name == "Оплачен":
            order.is_paid = True
            logger.info("Автоматически установлен флаг оплаты для заказа %s при изменении статуса на 'Оплачен'", order_id)
            session.add(order)
        
        # Получаем информацию о старом статусе
        old_status = await session.get(OrderStatusModel, order.status_id)
        old_status_name = old_status.name if old_status else "Неизвестный"
        
        # Фиксируем изменения в базе данных
        await session.commit()
        
        # Инвалидируем кэш заказа
        await invalidate_order_cache(order_id)
        # Инвалидируем кэш статистики
        await invalidate_statistics_cache()
        logger.info("Кэш заказа %s и статистики инвалидирован после изменения статуса с '%s' на '%s'", order_id, old_status_name, new_status_name)
        
        # Обновляем заказ в сессии
        updated_order = await get_order_by_id(session, order_id)

        # Если у заказа есть промокод, загружаем его ПЕРЕД отправкой уведомления
        if updated_order.promo_code_id:
            # Загружаем промокод
            promo_code = await session.get(PromoCodeModel, updated_order.promo_code_id)
            if promo_code:
                # Создаем словарь с данными промокода для передачи в RabbitMQ
                updated_order.promo_code_dict = {
                    "code": promo_code.code,
                    "discount_percent": promo_code.discount_percent or 0
                }
                logger.info("Для заказа %s загружен промокод %s (при обновлении статуса)", updated_order.id, promo_code.code)
        
        # Отправляем уведомление об изменении статуса через Notifications Service
        try:
            # Проверяем, принадлежит ли заказ авторизованному пользователю
            if updated_order.user_id:
                # Для заказов авторизованных пользователей всегда проверяем настройки уведомлений
                logger.info("Отправка уведомления об изменении статуса для авторизованного пользователя: %s", updated_order.user_id)
                await check_notification_settings(updated_order.user_id, "order.status_changed", updated_order.id)
            else:
                # Для заказов неавторизованных пользователей проверяем флаг согласия на уведомления
                if updated_order.receive_notifications and updated_order.email:
                    logger.info("Отправка уведомления об изменении статуса для неавторизованного пользователя на email: %s", updated_order.email)
                    # Используем специальный метод для неавторизованных пользователей
                    await check_notification_settings(None, "order.status_changed", updated_order.id)
                else:
                    logger.info("Уведомление об изменении статуса не отправлено: неавторизованный пользователь не дал согласие или не указал email")
            
            logger.info("Отправлено событие 'order.status_changed' для заказа %s", order_id)
        except (ConnectionError, TimeoutError, ValueError) as e:
            logger.error("Ошибка при отправке события изменения статуса: %s", e)
        
        return OrderResponse.model_validate(updated_order)
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Ошибка при обновлении статуса заказа: %s", str(e))
        await session.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Произошла ошибка при обновлении статуса заказа: {str(e)}",
        ) from e

@router.post("/{order_id}/items", response_model=OrderDetailResponseWithPromo)
async def update_order_items_endpoint(
    order_id: int,
    items_data: OrderItemsUpdate,
    session: AsyncSession = Depends(get_db),
    current_user: Dict[str, Any] = Depends(get_admin_user)
):
    """Обновление товаров в заказе (админский доступ)"""
    logger.info("Запрос на обновление элементов заказа %s: items_to_add=%s items_to_update=%s items_to_remove=%s", order_id, items_data.items_to_add, items_data.items_to_update, items_data.items_to_remove)
    
    try:
        # Получаем текущий статус заказа
        updated_order = await get_order_with_all_relations(session, order_id)
        if not updated_order:
            raise HTTPException(status_code=404, detail="Заказ не найден")
        
        # Проверка статуса заказа - запрещаем редактировать товары для определенных статусов
        non_editable_statuses = ["Отправлен", "Доставлен", "Отменен", "Оплачен"]
        if updated_order.status and updated_order.status.name in non_editable_statuses:
            raise HTTPException(
                status_code=400, 
                detail=f"Редактирование товаров невозможно для заказа в статусе '{updated_order.status.name}'"
            )
        
        # Проверка статуса оплаты - запрещаем редактировать товары для оплаченных заказов
        if updated_order.is_paid:
            raise HTTPException(
                status_code=400, 
                detail="Редактирование товаров невозможно: заказ уже оплачен"
            )
        
        # Подготавливаем данные для обновления
        items_to_add = [item.model_dump() for item in items_data.items_to_add] if items_data.items_to_add else []
        items_to_update = items_data.items_to_update if items_data.items_to_update else {}
        items_to_remove = items_data.items_to_remove if items_data.items_to_remove else []
        
        # Вызываем сервисную функцию
        updated_order = await update_order_items(
            order_id=updated_order.id,
            items_to_add=items_to_add,
            items_to_update=items_to_update,
            items_to_remove=items_to_remove,
            session=session,
            user_id=int(current_user.get("sub")) if current_user.get("sub") else None
        )
        
        if not updated_order:
            raise HTTPException(status_code=400, detail="Ошибка при обновлении элементов заказа")
        
        # Создаем ответ с данными заказа
        order_dict = {
            "id": updated_order.id,
            "user_id": updated_order.user_id,
            "status_id": updated_order.status_id,
            "status": {
                "id": updated_order.status.id,
                "name": updated_order.status.name,
                "description": updated_order.status.description,
                "color": updated_order.status.color,
                "allow_cancel": updated_order.status.allow_cancel,
                "is_final": updated_order.status.is_final,
                "sort_order": updated_order.status.sort_order
            },
            "created_at": updated_order.created_at,
            "updated_at": datetime.now(),
            "total_price": updated_order.total_price,
            "promo_code_id": updated_order.promo_code_id,
            "discount_amount": updated_order.discount_amount,
            "full_name": updated_order.full_name,
            "email": updated_order.email,
            "phone": updated_order.phone,
            "delivery_address": updated_order.delivery_address,
            "comment": updated_order.comment,
            "is_paid": updated_order.is_paid,
            "items": [
                {
                    "id": item.id,
                    "order_id": item.order_id,
                    "product_id": item.product_id,
                    "product_name": item.product_name,
                    "product_price": item.product_price,
                    "quantity": item.quantity,
                    "total_price": item.total_price
                } 
                for item in updated_order.items
            ],
            "order_number": f"{updated_order.id}-{updated_order.created_at.year}"
        }
        
        # Получаем промокод отдельным запросом если необходимо
        if updated_order.promo_code_id:
            promo_result = await session.execute(select(PromoCodeModel).where(PromoCodeModel.id == updated_order.promo_code_id))
            promo = promo_result.scalar_one_or_none()
            if promo:
                order_dict["promo_code"] = {
                    "id": promo.id,
                    "code": promo.code,
                    "discount_percent": promo.discount_percent,
                    "discount_amount": promo.discount_amount,
                    "valid_until": promo.valid_until,
                    "is_active": promo.is_active,
                    "created_at": promo.created_at,
                    "updated_at": datetime.now()
                }
        
        # Получаем информацию о доставке отдельным запросом
        delivery_query = select(DeliveryInfoModel).where(DeliveryInfoModel.order_id == updated_order.id)
        delivery_result = await session.execute(delivery_query)
        delivery_info = delivery_result.scalar_one_or_none()
        
        if delivery_info:
            order_dict["delivery_info"] = {
                "id": delivery_info.id,
                "order_id": delivery_info.order_id,
                "delivery_type": delivery_info.delivery_type,
                "boxberry_point_id": delivery_info.boxberry_point_id,
                "boxberry_point_address": delivery_info.boxberry_point_address,
                "delivery_cost": delivery_info.delivery_cost,
                "tracking_number": delivery_info.tracking_number,
                "label_url_boxberry": delivery_info.label_url_boxberry
            }
        
        # Получаем историю статусов отдельным запросом
        status_history_query = select(OrderStatusHistoryModel).where(
            OrderStatusHistoryModel.order_id == updated_order.id
        ).options(selectinload(OrderStatusHistoryModel.status))
        status_history_result = await session.execute(status_history_query)
        status_history_items = status_history_result.scalars().all()
        
        # Формируем историю статусов вручную
        order_dict["status_history"] = []
        for history in status_history_items:
            history_dict = {
                "id": history.id,
                "order_id": history.order_id,
                "status_id": history.status_id,
                "changed_at": history.changed_at,
                "changed_by_user_id": history.changed_by_user_id,
                "notes": history.notes,
            }
            
            if history.status:
                history_dict["status"] = {
                    "id": history.status.id,
                    "name": history.status.name,
                    "description": history.status.description,
                    "color": history.status.color,
                    "allow_cancel": history.status.allow_cancel,
                    "is_final": history.status.is_final,
                    "sort_order": history.status.sort_order
                }
            else:
                history_dict["status"] = None
                
            order_dict["status_history"].append(history_dict)
        
        # Получаем все необходимые поля из модели напрямую
        order_dict["personal_data_agreement"] = updated_order.personal_data_agreement
        order_dict["is_payment_on_delivery"] = getattr(updated_order, 'is_payment_on_delivery', True)
        order_dict["receive_notifications"] = getattr(updated_order, 'receive_notifications', None)
        
        # Создаем и возвращаем ответ из словаря вместо прямой валидации модели
        return OrderDetailResponseWithPromo.model_validate(order_dict)
    except (HTTPException, ValueError, AttributeError, KeyError) as e:
        logger.error("Ошибка при обновлении элементов заказа: %s", str(e))
        raise HTTPException(
            status_code=400,
            detail=f"Ошибка при обновлении элементов заказа: {str(e)}"
        )

@router.post("/batch-status", response_model=List[OrderResponse])
async def update_batch_status(
    update_data: BatchStatusUpdate,
    current_user: Dict[str, Any] = Depends(get_admin_user),
    session: AsyncSession = Depends(get_db)
):
    """
    Обновление статуса для нескольких заказов одновременно (только для администраторов).
    
    - **update_data**: Данные для массового обновления статусов
    """
    try:
        # Проверяем существование статуса
        order_status = await session.get(OrderStatusModel, update_data.status_id)
        if not order_status:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Статус с ID {update_data.status_id} не найден",
            )
        
        updated_orders = []
        for order_id in update_data.order_ids:
            try:
                # Создаем данные для обновления
                order_data = OrderUpdate(status_id=update_data.status_id)
                
                # Обновляем заказ
                order = await update_order(
                    session=session,
                    order_id=order_id,
                    order_data=order_data
                )
                
                if order:
                    # Добавляем запись в историю статусов
                    await OrderStatusHistoryModel.add_status_change(
                        session=session,
                        order_id=order.id,
                        status_id=update_data.status_id,
                        changed_by_user_id=current_user["user_id"],
                        notes=update_data.notes or "Массовое обновление статуса"
                    )
                    
                    loaded_order = await get_order_by_id(session, order.id)
                    updated_orders.append(loaded_order)
            except (ValueError, HTTPException, AttributeError) as e:
                logger.error("Ошибка при обновлении заказа %s: %s", order_id, str(e))
                # Продолжаем с другими заказами
        
        # Коммитим сессию
        await session.commit()
        
        # Преобразуем модели в схемы
        return [OrderResponse.model_validate(order) for order in updated_orders]
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Непредвиденная ошибка при массовом обновлении статусов: %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Произошла ошибка при массовом обновлении статусов",
        ) from e

@router.post("", response_model=OrderResponseWithPromo, status_code=status.HTTP_201_CREATED)
async def create_order_admin(
    order_data: AdminOrderCreate,
    current_user: Dict[str, Any] = Depends(get_admin_user),
    session: AsyncSession = Depends(get_db)
):
    """
    Создание нового заказа из админки. Доступно только для администраторов.
    
    - **order_data**: Данные для создания заказа
    """
    try:
        logger.info("Получен запрос на создание заказа из админки: %s", order_data)
        
        # Получаем токен авторизации администратора для запросов к другим сервисам
        token = current_user.get("token")
        
        # Если передан user_id, используем его, иначе заказ создается без привязки к пользователю
        user_id = order_data.user_id

        # Обработка пустого промокода
        promo_code = order_data.promo_code
        if promo_code == "" or (promo_code and len(promo_code) < 3):
            promo_code = None
            
        # Нормализация телефона
        phone = order_data.phone
        # Добавляем префикс, если его нет
        if phone and not (phone.startswith('8') or phone.startswith('+7')):
            phone = '8' + phone
        
        # Убеждаемся, что телефон имеет нужную длину
        if phone and len(phone) < 11:
            if phone.startswith('8'):
                # Дополняем до 11 цифр для формата 8XXXXXXXXXX
                missing_digits = 11 - len(phone)
                if missing_digits > 0:
                    phone = phone + '0' * missing_digits
            elif phone.startswith('+7'):
                # Дополняем до 12 цифр для формата +7XXXXXXXXXX
                missing_digits = 12 - len(phone)
                if missing_digits > 0:
                    phone = phone + '0' * missing_digits
                    
        # Проверяем тип доставки
        if not order_data.delivery_info.delivery_type:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Необходимо выбрать способ доставки",
            )
            
        # Для пунктов выдачи проверяем, что указан адрес пункта
        if order_data.delivery_info.delivery_type in ["boxberry_pickup_point", "cdek_pickup_point"] and not order_data.delivery_info.boxberry_point_address:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Необходимо указать адрес пункта выдачи",
            )
        
        # Создаем объект доставки для нового заказа
        delivery_info = order_data.delivery_info
        
        # Преобразуем данные из AdminOrderCreate в OrderCreate
        create_data = OrderCreate(
            items=order_data.items,
            full_name=order_data.full_name,
            email=order_data.email,
            phone=phone,  # Используем нормализованный телефон
            delivery_address=order_data.delivery_address,
            comment=order_data.comment,
            promo_code=promo_code,  # Используем обработанный промокод
            personal_data_agreement=True,  # Для админа всегда True
            delivery_info=delivery_info,
            is_payment_on_delivery=order_data.is_payment_on_delivery  # Добавляем флаг оплаты при получении
        )
        
        # Проверяем наличие товаров
        product_ids = [item.product_id for item in order_data.items]
        availability = await check_products_availability(product_ids)
        
        # Проверяем, все ли товары доступны
        unavailable_products = [pid for pid, available in availability.items() if not available]
        if unavailable_products:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Товары с ID {unavailable_products} недоступны для заказа",
            )
        
        # Создаем заказ через общую функцию
        order = await create_order(
            session=session,
            user_id=user_id,
            order_data=create_data,
            product_service_url=PRODUCT_SERVICE_URL,
            token=token
        )
        
        # Если указан начальный статус и он отличается от дефолтного
        if order_data.status_id:
            try:
                status_data = OrderStatusHistoryCreate(
                    status_id=order_data.status_id,
                    notes="Статус установлен при создании заказа администратором"
                )
                order = await change_order_status(
                    session=session,
                    order_id=order.id,
                    status_data=status_data,
                    user_id=current_user["user_id"],
                    is_admin=True
                )
                logger.info("Для нового заказа %s установлен статус %s", order.id, order_data.status_id)
            except (ValueError, HTTPException, AttributeError) as e:
                logger.warning("Не удалось установить начальный статус %s для заказа %s: %s", order_data.status_id, order.id, str(e))
        
        # Устанавливаем флаг оплаты, если указан
        if order_data.is_paid:
            order.is_paid = True
            await session.commit()
            logger.info("Для нового заказа %s установлен флаг оплаты", order.id)
        
        # Явно коммитим сессию, чтобы убедиться, что все связанные данные загружены
        await session.commit()
        
        # Вручную запрашиваем заказ со всеми связанными данными
        loaded_order = await get_order_by_id(session, order.id)
        
        # Если у заказа есть промокод, загружаем его
        if loaded_order.promo_code_id:
            # Загружаем промокод
            promo_code = await session.get(PromoCodeModel, loaded_order.promo_code_id)
            if promo_code:
                # Создаем словарь с данными промокода для передачи в RabbitMQ
                loaded_order.promo_code_dict = {
                    "code": promo_code.code,
                    "discount_percent": promo_code.discount_percent or 0
                }
                logger.info("Для нового заказа %s загружен промокод %s", loaded_order.id, promo_code.code)
        
        # Отправляем подтверждение заказа на email через Notifications Service
        try:
            if order_data.email and user_id != None:
                logger.info("Отправка подтверждения заказа на email: %s", order_data.email)
                await check_notification_settings(loaded_order.user_id, "order.created", loaded_order.id)
        except (ConnectionError, TimeoutError, ValueError) as e:
            logger.error("Ошибка при отправке email о заказе: %s", str(e))
        
        # Явно инвалидируем кэш заказов перед возвратом ответа
        await invalidate_order_cache(order.id)
        await invalidate_statistics_cache()  # Также инвалидируем кэш статистики и отчетов
        logger.info("Кэш заказа %s и связанных списков инвалидирован перед возвратом ответа", order.id)
        
        # Преобразуем модель в схему без промокода
        order_response = OrderResponse.model_validate(loaded_order)
        
        # Создаем расширенный ответ с возможностью добавления промокода
        response_with_promo = OrderResponseWithPromo(**order_response.model_dump())
        
        # Если у заказа есть промокод, загружаем его данные вручную
        if loaded_order.promo_code_id:
            try:
                promo_code = await session.get(PromoCodeModel, loaded_order.promo_code_id)
                if promo_code:
                    response_with_promo.promo_code = PromoCodeResponse.model_validate(promo_code)
                    logger.info("Для нового заказа %s загружен промокод %s", loaded_order.id, promo_code.code)
            except (ValueError, AttributeError, KeyError) as e:
                logger.warning("Не удалось загрузить промокод %s для заказа %s: %s", loaded_order.promo_code_id, loaded_order.id, str(e))
        
        return response_with_promo
    except ValueError as e:
        logger.error("Ошибка при создании заказа из админки: %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(e),
        ) from e
    except Exception as e:
        logger.error("Непредвиденная ошибка при создании заказа из админки: %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Произошла ошибка при создании заказа",
        ) from e 

@router.get("/reports/download", dependencies=[Depends(get_admin_user)])
async def generate_orders_report(
    format: str = Query(..., description="Формат отчета: csv, excel, pdf, word"),
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
    session: AsyncSession = Depends(get_db)
):
    """
    Генерация отчета по заказам в различных форматах
    
    - **format**: Формат отчета (csv, excel, pdf, word)
    - **date_from**: Дата начала периода (YYYY-MM-DD)
    - **date_to**: Дата окончания периода (YYYY-MM-DD)
    """
    try:
        # Получаем статистику за указанный период
        statistics = await get_order_statistics_by_date(session, date_from, date_to)
        
        # Создаем датафрейм со статистикой
        stats_data = {
            "Метрика": ["Всего заказов", "Общая сумма заказов", "Средняя стоимость заказа", "Сумма отмененных заказов"],
            "Значение": [
                statistics.total_orders,
                statistics.total_revenue,
                round(statistics.average_order_value, 2),
                statistics.canceled_orders_revenue
            ]
        }
        
        # Добавляем статистику по статусам
        for status_name, count in statistics.orders_by_status.items():
            stats_data["Метрика"].append(f"Заказов со статусом '{status_name}'")
            stats_data["Значение"].append(count)
        
        df_stats = pd.DataFrame(stats_data)
        
        # Получаем список заказов за указанный период
        # Создаем фильтры для запроса заказов
        filters = OrderFilterParams(
            page=1,
            size=100,  # Максимально допустимое значение
            date_from=date_from,
            date_to=date_to
        )
        
        # Получаем заказы
        orders, _ = await get_orders(session, filters)
        
        # Создаем датафрейм с заказами
        orders_data = []
        for order in orders:
            try:
                order_number = order.order_number if hasattr(order, 'order_number') else f"{order.id}-{order.created_at.year}"
                created_at_str = order.created_at.strftime("%Y-%m-%d %H:%M") if order.created_at else "Н/Д"
                status_name = order.status.name if order.status else "Неизвестный"
                is_paid_str = "Да" if order.is_paid else "Нет"
                
                orders_data.append([
                    str(order.id),
                    order_number,
                    created_at_str,
                    status_name,
                    f"{order.total_price:.2f} руб.",
                    is_paid_str
                ])
            except Exception as e:
                logger.error(f"Ошибка при обработке заказа {order.id} для PDF: {str(e)}")
                # Добавляем строку с минимальной информацией, чтобы не прерывать генерацию отчета
                orders_data.append([
                    str(order.id),
                    f"{order.id}-????",
                    "Ошибка даты",
                    "Ошибка статуса",
                    f"{order.total_price:.2f} руб.",
                    "Ошибка"
                ])
        
        df_orders = pd.DataFrame(orders_data)
        
        # Период для имени файла
        period = ""
        if date_from and date_to:
            period = f"_{date_from}_{date_to}"
        elif date_from:
            period = f"_{date_from}"
        elif date_to:
            period = f"_{date_to}"
        
        # Обработка в зависимости от формата
        if format.lower() == "csv":
            # Создаем буфер для CSV файла
            output = io.StringIO()
            
            # Записываем заголовок
            output.write("Отчет по заказам\n")
            if date_from and date_to:
                output.write(f"Период: с {date_from} по {date_to}\n\n")
            elif date_from:
                output.write(f"Период: с {date_from}\n\n")
            elif date_to:
                output.write(f"Период: по {date_to}\n\n")
            else:
                output.write("Период: все время\n\n")
            
            # Записываем статистику
            output.write("Общая статистика:\n")
            df_stats.to_csv(output, index=False, sep=';')
            
            output.write("\n\nСписок заказов:\n")
            df_orders.to_csv(output, index=False, sep=';')
            
            # Возвращаем StreamingResponse
            output.seek(0)
            
            # Форматируем имя файла                
            filename = f"orders_report{period}.csv"
            
            return StreamingResponse(
                iter([output.getvalue()]),
                media_type="text/csv",
                headers={"Content-Disposition": f"attachment; filename={filename}"}
            )
        
        elif format.lower() == "excel":
            # Создаем буфер для Excel файла
            output = io.BytesIO()
            
            # Создаем Excel файл
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                # Добавляем лист со статистикой
                df_stats.to_excel(writer, sheet_name='Статистика', index=False)
                
                # Добавляем лист с заказами
                df_orders.to_excel(writer, sheet_name='Заказы', index=False)
            
            # Возвращаем StreamingResponse
            output.seek(0)
            
            # Форматируем имя файла с правильным расширением
            filename = f"orders_report{period}.xlsx"
            
            return StreamingResponse(
                io.BytesIO(output.getvalue()),
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                headers={"Content-Disposition": f"attachment; filename={filename}"}
            )
        
        elif format.lower() == "pdf":
            try:
                
                buffer = io.BytesIO()
                
                # Регистрируем системный шрифт DejaVu для поддержки кириллицы
                try:
                    # В Ubuntu/Debian системах DejaVu расположен тут
                    dejavu_sans_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
                    dejavu_serif_path = "/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf"
                    dejavu_sans_bold_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"
                    
                    if os.path.exists(dejavu_sans_path) and os.path.exists(dejavu_serif_path):
                        # Регистрируем шрифты
                        pdfmetrics.registerFont(TTFont('DejaVuSans', dejavu_sans_path))
                        pdfmetrics.registerFont(TTFont('DejaVuSerif', dejavu_serif_path))
                        
                        # Проверяем наличие жирного шрифта
                        if os.path.exists(dejavu_sans_bold_path):
                            pdfmetrics.registerFont(TTFont('DejaVuSans-Bold', dejavu_sans_bold_path))
                            font_name_bold = 'DejaVuSans-Bold'
                        else:
                            font_name_bold = 'DejaVuSans'
                        
                        font_name = 'DejaVuSans'
                        logger.info("Шрифт DejaVu успешно зарегистрирован для PDF отчетов")
                    else:
                        # Если шрифтов нет, используем стандартные 
                        font_name = 'Helvetica'
                        font_name_bold = 'Helvetica-Bold'
                        logger.warning("Шрифты DejaVu не найдены, используем стандартные шрифты")
                except Exception as e:
                    # В случае ошибки используем стандартные шрифты
                    logger.error(f"Ошибка при регистрации шрифта: {e}")
                    font_name = 'Helvetica'
                    font_name_bold = 'Helvetica-Bold'
                
                # Создаем PDF документ с указанием кодировки UTF-8
                doc = SimpleDocTemplate(buffer, pagesize=landscape(letter), encoding='utf-8')
                elements = []
                
                # Настраиваем стили с использованием DejaVu шрифтов
                styles = getSampleStyleSheet()
                
                # Создаем собственные стили с кодировкой UTF-8 для поддержки кириллицы
                title_style = ParagraphStyle(
                    'CustomTitle',
                    parent=styles['Heading1'],
                    alignment=TA_CENTER,
                    fontName=font_name_bold,
                    fontSize=16,
                    spaceAfter=12,
                    encoding='utf-8'
                )
                
                subtitle_style = ParagraphStyle(
                    'CustomSubtitle',
                    parent=styles['Heading2'],
                    fontName=font_name_bold,
                    fontSize=14,
                    spaceAfter=6,
                    encoding='utf-8'
                )
                
                normal_style = ParagraphStyle(
                    'CustomNormal',
                    parent=styles['Normal'],
                    fontName=font_name,
                    fontSize=10,
                    encoding='utf-8'
                )
                
                # Добавляем заголовок
                elements.append(Paragraph("Отчет по заказам", title_style))
                elements.append(Spacer(1, 10))
                
                # Период отчета
                if date_from and date_to:
                    elements.append(Paragraph(f"Период: с {date_from} по {date_to}", subtitle_style))
                elif date_from:
                    elements.append(Paragraph(f"Период: с {date_from}", subtitle_style))
                elif date_to:
                    elements.append(Paragraph(f"Период: по {date_to}", subtitle_style))
                else:
                    elements.append(Paragraph("Период: все время", subtitle_style))
                
                elements.append(Spacer(1, 15))
                
                # Общая статистика
                elements.append(Paragraph("Общая статистика", subtitle_style))
                elements.append(Spacer(1, 5))
                
                # Создаем данные для таблицы статистики
                stats_data = [['Метрика', 'Значение']]
                stats_data.append(['Всего заказов', str(statistics.total_orders)])
                stats_data.append(['Общая сумма заказов', f"{statistics.total_revenue:.2f} руб."])
                stats_data.append(['Средняя стоимость заказа', f"{statistics.average_order_value:.2f} руб."])
                stats_data.append(['Сумма отмененных заказов', f"{statistics.canceled_orders_revenue:.2f} руб."])
                
                # Добавляем статистику по статусам
                for status_name, count in statistics.orders_by_status.items():
                    stats_data.append([f"Заказов со статусом '{status_name}'", str(count)])
                
                # Создаем таблицу с чистыми, светлыми цветами
                stats_table = Table(stats_data, colWidths=[300, 150])
                
                # Создаем базовые стили таблицы статистики
                stats_table_styles = [
                    # Светло-голубой фон для заголовка
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E6F3FF')),
                    ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                    ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), font_name_bold),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('ALIGN', (1, 1), (1, -1), 'RIGHT'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('FONTNAME', (0, 1), (-1, -1), font_name),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('TOPPADDING', (0, 0), (-1, -1), 6),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ]
                
                # Добавляем фон для четных строк только если есть достаточно данных
                stats_row_count = len(stats_data)
                if stats_row_count > 2:  # Если есть хотя бы одна строка данных помимо заголовка
                    # Добавляем светло-серый фон для четных строк
                    stats_table_styles.append(('BACKGROUND', (0, 2), (-1, 2), colors.HexColor('#F5F5F5')))
                    
                    if stats_row_count > 4:  # Если есть хотя бы две строки данных
                        stats_table_styles.append(('BACKGROUND', (0, 4), (-1, 4), colors.HexColor('#F5F5F5')))
                        
                        if stats_row_count > 6:  # Если есть хотя бы три строки данных
                            stats_table_styles.append(('BACKGROUND', (0, 6), (-1, 6), colors.HexColor('#F5F5F5')))
                
                stats_table.setStyle(TableStyle(stats_table_styles))
                
                elements.append(stats_table)
                elements.append(Spacer(1, 20))
                
                # Список заказов
                elements.append(Paragraph("Список заказов", subtitle_style))
                elements.append(Spacer(1, 5))
                
                # Создаем таблицу со списком заказов
                if orders:
                    # Заголовки таблицы
                    orders_data = [['ID', 'Номер заказа', 'Дата создания', 'Статус', 'Сумма', 'Оплачен']]
                    
                    # Данные заказов
                    for order in orders:
                        try:
                            order_number = order.order_number if hasattr(order, 'order_number') else f"{order.id}-{order.created_at.year}"
                            orders_data.append([
                                str(order.id),
                                order_number,
                                order.created_at.strftime("%Y-%m-%d %H:%M"),
                                order.status.name if order.status else "Неизвестный",
                                f"{order.total_price:.2f} руб.",
                                "Да" if order.is_paid else "Нет"
                            ])
                        except Exception as e:
                            logger.error(f"Ошибка при обработке заказа {order.id} для PDF: {str(e)}")
                            # Добавляем строку с минимальной информацией, чтобы не прерывать генерацию отчета
                            orders_data.append([
                                str(order.id),
                                f"{order.id}-????",
                                "Ошибка даты",
                                "Ошибка статуса",
                                f"{order.total_price:.2f} руб.",
                                "Ошибка"
                            ])
                    
                    # Создаем таблицу с чистыми, светлыми цветами
                    orders_table = Table(orders_data, colWidths=[35, 70, 120, 100, 150, 70])
                    
                    # Создаем базовые стили таблицы
                    table_styles = [
                        # Светло-голубой фон для заголовка
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E6F3FF')),
                        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), font_name_bold),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                        ('ALIGN', (0, 0), (0, -1), 'CENTER'),  # ID по центру
                        ('ALIGN', (4, 1), (4, -1), 'RIGHT'),   # Суммы по правому краю (теперь индекс 4 вместо 5)
                        ('ALIGN', (5, 1), (5, -1), 'CENTER'),  # Оплачен по центру (теперь индекс 5 вместо 6)
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        ('FONTNAME', (0, 1), (-1, -1), font_name),
                        ('FONTSIZE', (0, 0), (-1, -1), 9),     # Уменьшенный размер шрифта для большого количества данных
                        ('TOPPADDING', (0, 0), (-1, -1), 6),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ]
                    
                    # Добавляем фон для четных строк только если есть достаточно данных
                    row_count = len(orders_data)
                    if row_count > 2:  # Если есть хотя бы одна строка данных помимо заголовка
                        # Добавляем светло-серый фон для четных строк (строка с индексом 2 - это первая четная строка данных)
                        table_styles.append(('BACKGROUND', (0, 2), (-1, 2), colors.HexColor('#F5F5F5')))
                        
                        if row_count > 4:  # Если есть хотя бы две строки данных
                            table_styles.append(('BACKGROUND', (0, 4), (-1, 4), colors.HexColor('#F5F5F5')))
                            
                            if row_count > 6:  # Если есть хотя бы три строки данных
                                table_styles.append(('BACKGROUND', (0, 6), (-1, 6), colors.HexColor('#F5F5F5')))
                    
                    orders_table.setStyle(TableStyle(table_styles))
                    
                    elements.append(orders_table)
                else:
                    elements.append(Paragraph("Нет заказов за выбранный период", normal_style))
                
                # Информация о дате формирования отчета
                elements.append(Spacer(1, 20))
                elements.append(Paragraph(f"Отчет сформирован: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", normal_style))
                
                # Строим PDF документ
                doc.build(elements)
                
                # Готовим ответ
                buffer.seek(0)
                
                # Имя файла
                filename = f"orders_report{period}.pdf"
                
                return StreamingResponse(
                    buffer,
                    media_type="application/pdf",
                    headers={"Content-Disposition": f"attachment; filename={filename}"}
                )
            except Exception as e:
                logger.error(f"Ошибка при создании PDF отчета: {str(e)}")
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"Ошибка при создании PDF отчета: {str(e)}"
                )
        
        elif format.lower() == "word":
            try:
                # Импортируем необходимые библиотеки
                from docx import Document
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                # Создаем документ Word
                doc = Document()
                
                # Заголовок отчета
                title = doc.add_heading('Отчет по заказам', level=1)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Период отчета
                if date_from and date_to:
                    doc.add_heading(f"Период: с {date_from} по {date_to}", level=2)
                elif date_from:
                    doc.add_heading(f"Период: с {date_from}", level=2)
                elif date_to:
                    doc.add_heading(f"Период: по {date_to}", level=2)
                else:
                    doc.add_heading("Период: все время", level=2)
                
                # Добавляем раздел общей статистики
                doc.add_heading('Общая статистика', level=2)
                
                # Таблица со статистикой
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                
                # Заголовки таблицы
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Метрика'
                hdr_cells[1].text = 'Значение'
                
                # Форматирование заголовков
                for cell in table.rows[0].cells:
                    cell.paragraphs[0].runs[0].bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Данные статистики
                row = table.add_row().cells
                row[0].text = 'Всего заказов'
                row[1].text = str(statistics.total_orders)
                
                row = table.add_row().cells
                row[0].text = 'Общая сумма заказов'
                row[1].text = f"{statistics.total_revenue:.2f} руб."
                
                row = table.add_row().cells
                row[0].text = 'Средняя стоимость заказа'
                row[1].text = f"{statistics.average_order_value:.2f} руб."
                
                row = table.add_row().cells
                row[0].text = 'Сумма отмененных заказов'
                row[1].text = f"{statistics.canceled_orders_revenue:.2f} руб."
                
                # Статистика по статусам
                for status_name, count in statistics.orders_by_status.items():
                    row = table.add_row().cells
                    row[0].text = f"Заказов со статусом '{status_name}'"
                    row[1].text = str(count)
                
                # Добавляем раздел списка заказов
                doc.add_heading('Список заказов', level=2)
                
                # Таблица заказов
                if orders:
                    table = doc.add_table(rows=1, cols=6)
                    table.style = 'Table Grid'
                    
                    # Заголовки таблицы
                    hdr_cells = table.rows[0].cells
                    headers = ['ID', 'Номер заказа', 'Дата создания', 'Статус', 'Сумма', 'Оплачен']
                    for i, header in enumerate(headers):
                        hdr_cells[i].text = header
                    
                    # Форматирование заголовков
                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].bold = True
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Данные заказов
                    for order in orders:
                        try:
                            order_number = order.order_number if hasattr(order, 'order_number') else f"{order.id}-{order.created_at.year}"
                            row_cells = table.add_row().cells
                            row_cells[0].text = str(order.id)
                            row_cells[1].text = order_number
                            row_cells[2].text = order.created_at.strftime("%Y-%m-%d %H:%M")
                            row_cells[3].text = order.status.name if order.status else "Неизвестный"
                            row_cells[4].text = f"{order.total_price:.2f} руб."
                            row_cells[5].text = "Да" if order.is_paid else "Нет"
                        except Exception as e:
                            logger.error(f"Ошибка при обработке заказа {order.id} для Word отчета: {str(e)}")
                            # Добавляем строку с минимальной информацией, чтобы не прерывать генерацию отчета
                            row_cells = table.add_row().cells
                            row_cells[0].text = str(order.id) if hasattr(order, 'id') else "Н/Д"
                            row_cells[1].text = f"{order.id}-????" if hasattr(order, 'id') else "Н/Д"
                            row_cells[2].text = "Н/Д"
                            row_cells[3].text = "Н/Д"
                            row_cells[4].text = f"{order.total_price:.2f} руб." if hasattr(order, 'total_price') else "0.00 руб."
                            row_cells[5].text = "Н/Д"
                else:
                    doc.add_paragraph("Нет заказов за выбранный период")
                
                # Сохраняем документ в буфер
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                # Имя файла с правильным расширением
                filename = f"orders_report{period}.docx"
                
                return StreamingResponse(
                    buffer,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": f"attachment; filename={filename}"}
                )
            except Exception as e:
                logger.error(f"Ошибка при создании Word отчета: {str(e)}")
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"Ошибка при создании Word отчета: {str(e)}"
                )
        
        else:
            # Возвращаем ошибку, если формат не поддерживается
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Неподдерживаемый формат отчета: {format}. Поддерживаемые форматы: csv, excel, pdf, word"
            )
    
    except Exception as e:
        logger.error("Ошибка при генерации отчета: %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Произошла ошибка при генерации отчета"
        ) from e


@router.post("/{order_id}/payment-status", response_model=OrderResponse,dependencies=[Depends(get_admin_user)])
async def update_order_payment_status(
    order_id: int = Path(..., ge=1),
    payment_data: dict = Body(...),
    session: AsyncSession = Depends(get_db)
):
    """
    Обновить статус оплаты заказа (is_paid) отдельно от других изменений.
    """
    try:
        order = await get_order_by_id(session, order_id)
        if not order:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=f"Заказ с ID {order_id} не найден")
        is_paid = payment_data.get("is_paid")
        if is_paid is None:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Поле is_paid обязательно")
        # Можно добавить бизнес-валидацию, если нужно
        order.is_paid = bool(is_paid)
        await session.commit()
        await invalidate_order_cache(order_id)
        updated_order = await get_order_by_id(session, order_id)
        return OrderResponse.model_validate(updated_order)
    except HTTPException:
        raise
    except Exception as e:
        logger.exception(f"Ошибка при обновлении статуса оплаты заказа: {str(e)}")
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Ошибка при обновлении оплаты: {str(e)}")