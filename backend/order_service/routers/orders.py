"""API эндпоинты сервиса заказов для управления заказами, включая создание, получение, обновление и управление статусами.
Предоставляет интерфейсы как для пользователей, так и для администраторов."""

import os
import logging
from typing import Optional, List, Dict, Any
from datetime import datetime
import hashlib
from dotenv import load_dotenv
from fastapi import APIRouter, Depends, HTTPException, Query, Path, status, Body
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import text
from notification_api import check_notification_settings
import pandas as pd
import io
import csv
import openpyxl

from database import get_db
from models import OrderStatusModel, OrderStatusHistoryModel, PromoCodeModel
from schemas import (
    OrderCreate, OrderUpdate, OrderResponse, OrderDetailResponse, 
    OrderStatusHistoryCreate, PaginatedResponse, OrderStatistics, BatchStatusUpdate,
    OrderItemsUpdate, OrderItemsUpdateResponse, PromoCodeResponse, OrderResponseWithPromo, OrderDetailResponseWithPromo,
    AdminOrderCreate, OrderFilterParams
)
from services import (
    create_order, get_order_by_id, get_orders, update_order, 
    change_order_status, cancel_order, get_order_statistics, get_user_order_statistics,
    update_order_items, get_order_statistics_by_date
)
from dependencies import (
    get_current_user, get_admin_user, get_order_filter_params,
    check_products_availability, get_products_info, verify_service_jwt
)
from cache import (
    get_cached_order, cache_order, invalidate_order_cache,
    get_cached_user_orders, cache_user_orders,
    get_cached_order_statistics, cache_order_statistics, invalidate_statistics_cache,
    get_cached_orders_list, cache_orders_list, CacheKeys, get_cached_data, set_cached_data,
    DEFAULT_CACHE_TTL
)

# Загрузка переменных окружения
load_dotenv()

# Настройка логирования
logger = logging.getLogger("order_router")

# Получение URL сервиса продуктов
PRODUCT_SERVICE_URL = os.getenv("PRODUCT_SERVICE_URL", "http://localhost:8001")

# Создание роутера
router = APIRouter(
    prefix="/orders",
    tags=["orders"],
    responses={404: {"description": "Not found"}},
)

# Создание роутера для админ-панели
admin_router = APIRouter(
    prefix="/admin/orders",
    tags=["admin_orders"],
    responses={404: {"description": "Not found"}},
)

@router.post("", response_model=OrderResponseWithPromo, status_code=status.HTTP_201_CREATED)
async def create_new_order(
    order_data: OrderCreate,
    current_user: Optional[Dict[str, Any]] = Depends(get_current_user),
    session: AsyncSession = Depends(get_db)
):
    """
    Создание нового заказа. Доступно как для авторизованных, так и для анонимных пользователей.
    
    - **order_data**: Данные для создания заказа
    """
    try:
        logger.info("Получен запрос на создание заказа: %s", order_data)
        
        # Получаем ID пользователя из токена (если пользователь авторизован)
        user_id = current_user.get("user_id") if current_user else None
        
        # Получаем токен авторизации (если пользователь авторизован)
        token = current_user.get("token") if current_user else None
        
        # Проверяем наличие товаров
        product_ids = [item.product_id for item in order_data.items]
        availability = await check_products_availability(product_ids)
        
        # Проверяем, все ли товары доступны
        unavailable_products = [pid for pid, available in availability.items() if not available]
        if unavailable_products:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Товары с ID {unavailable_products} недоступны для заказа",
            )
        
        # Создаем заказ
        order = await create_order(
            session=session,
            user_id=user_id,
            order_data=order_data,
            product_service_url=PRODUCT_SERVICE_URL,
            token=token
        )
        
        # Если пользователь авторизован, очищаем его корзину
        # if user_id:
        #     try:
        #         await clear_user_cart(user_id)
        #     except Exception as e:
        #         logger.warning(f"Не удалось очистить корзину пользователя: {str(e)}")
        
        # Явно коммитим сессию, чтобы убедиться, что все связанные данные загружены
        await session.commit()
        
        # Вручную запрашиваем заказ со всеми связанными данными
        loaded_order = await get_order_by_id(session, order.id)
        
        # Если у заказа есть промокод, загружаем его
        if loaded_order.promo_code_id:
            # Загружаем промокод
            promo_code = await session.get(PromoCodeModel, loaded_order.promo_code_id)
            if promo_code:
                # Создаем словарь с данными промокода для передачи в RabbitMQ
                loaded_order.promo_code_dict = {
                    "code": promo_code.code,
                    "discount_percent": promo_code.discount_percent or 0
                }
                logger.info("Для нового заказа %s загружен промокод %s", loaded_order.id, promo_code.code)
        
        # Отправляем подтверждение заказа на email через Notifications Service
        try:
            # Проверяем, авторизован ли пользователь
            if user_id:
                # Для авторизованных пользователей всегда проверяем настройки уведомлений
                logger.info("Отправка уведомления о создании заказа для авторизованного пользователя: %s", user_id)
                await check_notification_settings(loaded_order.user_id, "order.created", loaded_order.id)
            else:
                # Для неавторизованных пользователей проверяем согласие на получение уведомлений
                if loaded_order.receive_notifications and loaded_order.email:
                    logger.info("Отправка уведомления о создании заказа для неавторизованного пользователя на email: %s", order_data.email)
                    # Для неавторизованных пользователей отправляем уведомление напрямую, без проверки настроек
                    # Здесь можно использовать другой метод для отправки, который не проверяет настройки пользователя
                    # Например, direct_notification_service или аналогичный
                    await check_notification_settings(None, "order.created", loaded_order.id)
                else:
                    logger.info("Уведомление о создании заказа не отправлено: неавторизованный пользователь не дал согласие или не указал email")
        except (ConnectionError, TimeoutError, ValueError) as e:
            logger.error("Ошибка при отправке email о заказе: %s", str(e))
        
        # Явно инвалидируем кэш заказов перед возвратом ответа
        await invalidate_order_cache(order.id)
        logger.info("Кэш заказа %s и связанных списков инвалидирован перед возвратом ответа", order.id)
        
        # Преобразуем модель в схему без промокода
        order_response = OrderResponse.model_validate(loaded_order)
        
        # Создаем расширенный ответ с возможностью добавления промокода
        response_with_promo = OrderResponseWithPromo(**order_response.model_dump())
        
        # Если у заказа есть промокод, загружаем его данные вручную
        if loaded_order.promo_code_id:
            try:
                promo_code = await session.get(PromoCodeModel, loaded_order.promo_code_id)
                if promo_code:
                    response_with_promo.promo_code = PromoCodeResponse.model_validate(promo_code)
                    logger.info("Для нового заказа %s загружен промокод %s", loaded_order.id, promo_code.code)
            except (ValueError, AttributeError, KeyError) as e:
                logger.warning("Не удалось загрузить промокод %s для заказа %s: %s", loaded_order.promo_code_id, loaded_order.id, str(e))
        
        # Кэшируем результат в любом случае, но с разными ключами
        cache_key = f"{CacheKeys.ORDER_PREFIX}{order.id}"
        if user_id:
            # Для авторизованных пользователей - отдельный кэш
            cache_key = f"{CacheKeys.ORDER_PREFIX}{order.id}:user:{user_id}"
        
        await cache_order(order.id, response_with_promo, cache_key=cache_key)
        logger.info("Данные о заказе %s добавлены в кэш с ключом %s", order.id, cache_key)
        
        return response_with_promo
    except ValueError as e:
        logger.error("Ошибка при создании заказа: %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(e),
        ) from e
    except Exception as e:
        logger.error("Непредвиденная ошибка при создании заказа: %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Произошла ошибка при создании заказа",
        ) from e

@router.get("", response_model=PaginatedResponse)
async def list_my_orders(
    page: int = Query(1, ge=1),
    size: int = Query(10, ge=1, le=100),
    status_id: Optional[int] = None,
    current_user: Optional[Dict[str, Any]] = Depends(get_current_user),
    session: AsyncSession = Depends(get_db)
):
    """
    Получение списка заказов текущего пользователя.
    
    - **page**: Номер страницы
    - **size**: Размер страницы
    - **status_id**: ID статуса заказа
    """
    # Проверяем, авторизован ли пользователь
    if not current_user:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Для просмотра заказов необходима авторизация",
            headers={"WWW-Authenticate": "Bearer"},
        )
        
    try:
        # Получаем ID пользователя из токена
        user_id = current_user["user_id"]
        
        # Создаем параметры фильтрации
        filters = get_order_filter_params(
            page=page,
            size=size,
            status_id=status_id,
            user_id=user_id
        )
        
        # Создаем ключ для кэша на основе параметров фильтрации
        cache_key = f"p{page}_s{size}_st{status_id or 'all'}"
        
        # Пытаемся получить данные из кэша
        cached_orders = await get_cached_user_orders(user_id, cache_key)
        if cached_orders:
            logger.info("Данные о заказах пользователя %s получены из кэша", user_id)
            return cached_orders
        
        # Если данных нет в кэше, получаем из БД
        orders, total = await get_orders(
            session=session,
            filters=filters,
            user_id=user_id
        )
        
        # Формируем ответ с преобразованием моделей в схемы
        order_responses = []
        for order in orders:
            # Создаем базовый объект OrderResponse без промокода
            order_response = OrderResponse.model_validate(order)
            
            # Если у заказа есть промокод, создаем OrderResponseWithPromo
            if order.promo_code_id:
                # Создаем расширенный ответ с информацией о промокоде
                with_promo = OrderResponseWithPromo(**order_response.model_dump())
                
                try:
                    # Получаем промокод и устанавливаем его
                    promo_code = await session.get(PromoCodeModel, order.promo_code_id)
                    if promo_code:
                        with_promo.promo_code = PromoCodeResponse.model_validate(promo_code)
                        logger.info("Для заказа %s загружен промокод %s", order.id, promo_code.code)
                except (ValueError, AttributeError, KeyError) as e:
                    logger.warning("Не удалось загрузить промокод %s для заказа %s: %s", order.promo_code_id, order.id, str(e))
                
                order_responses.append(with_promo)
            else:
                # Заказ без промокода, используем обычный OrderResponse
                order_responses.append(order_response)
        
        response = PaginatedResponse(
            items=order_responses,
            total=total,
            page=filters.page,
            size=filters.size,
            pages=0  # Будет вычислено в валидаторе
        )
        
        # Кэшируем результат
        await cache_user_orders(user_id, cache_key, response)
        logger.info("Данные о заказах пользователя %s добавлены в кэш", user_id)
        
        return response
    except Exception as e:
        logger.error("Ошибка при получении списка заказов пользователя: %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Произошла ошибка при получении списка заказов",
        ) from e

@router.get("/statistics", response_model=OrderStatistics)
async def get_user_statistics(
    current_user: Optional[Dict[str, Any]] = Depends(get_current_user),
    session: AsyncSession = Depends(get_db)
):
    """
    Получение статистики по заказам текущего пользователя
    """
    # Проверяем, авторизован ли пользователь
    if not current_user:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Для получения статистики необходима авторизация",
            headers={"WWW-Authenticate": "Bearer"},
        )
        
    user_id = current_user.get("user_id")
    if not user_id:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Не удалось определить пользователя"
        )
    
    try:
        # Пытаемся получить данные из кэша
        cached_statistics = await get_cached_order_statistics(user_id)
        if cached_statistics:
            logger.info("Статистика заказов пользователя %s получена из кэша", user_id)
            return cached_statistics
        
        # Если данных нет в кэше, получаем из БД
        statistics = await get_user_order_statistics(session, user_id)
        
        # Кэшируем результат
        await cache_order_statistics(statistics, user_id)
        logger.info("Статистика заказов пользователя %s добавлена в кэш", user_id)
        
        return statistics
    except Exception as e:
        logger.error("Ошибка при получении статистики заказов пользователя: %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Ошибка при получении статистики заказов"
        ) from e

@router.get("/{order_id}", response_model=OrderDetailResponseWithPromo)
async def get_order(
    order_id: int = Path(..., ge=1),
    current_user: Optional[Dict[str, Any]] = Depends(get_current_user),
    session: AsyncSession = Depends(get_db)
):
    """
    Получение информации о заказе пользователя по ID.
    
    - **order_id**: ID заказа
    """
    try:
        # Проверяем, авторизован ли пользователь
        if not current_user:
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Для просмотра заказа необходима авторизация",
                headers={"WWW-Authenticate": "Bearer"},
            )
        
        # Определяем пользователя
        user_id = current_user["user_id"]
        
        # Пытаемся получить данные из кэша
        cached_order = await get_cached_order(order_id, user_id)
        if cached_order:
            logger.info("Данные о заказе %s получены из кэша для пользователя %s", order_id, user_id)
            return cached_order
            
        # Получаем заказ
        order = await get_order_by_id(
            session=session,
            order_id=order_id,
            user_id=user_id
        )
        
        if not order:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Заказ с ID {order_id} не найден",
            )
        
        # Преобразуем модель в схему без промокода
        order_response = OrderDetailResponse.model_validate(order)
        
        # Создаем расширенный ответ с возможностью добавления промокода
        response_with_promo = OrderDetailResponseWithPromo(**order_response.model_dump())
        
        # Вручную обрабатываем промокод
        if order.promo_code_id:
            try:
                promo_code = await session.get(PromoCodeModel, order.promo_code_id)
                if promo_code:
                    response_with_promo.promo_code = PromoCodeResponse.model_validate(promo_code)
                    logger.info("Для заказа %s загружен промокод %s", order.id, promo_code.code)
            except (ValueError, AttributeError, KeyError) as e:
                logger.warning("Не удалось загрузить промокод %s для заказа %s: %s", order.promo_code_id, order.id, str(e))
        
        # Кэшируем результат
        cache_key = f"{CacheKeys.ORDER_PREFIX}{order_id}"
        if user_id:
            cache_key = f"{CacheKeys.ORDER_PREFIX}{order_id}:user:{user_id}"
            
        await cache_order(order_id, response_with_promo, cache_key=cache_key)
        logger.info("Данные о заказе %s добавлены в кэш с ключом %s", order_id, cache_key)
        
        return response_with_promo
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Ошибка при получении информации о заказе: %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Произошла ошибка при получении информации о заказе",
        ) from e

@router.post("/{order_id}/cancel", response_model=OrderResponse)
async def cancel_order_endpoint(
    order_id: int = Path(..., ge=1),
    cancel_reason: Optional[str] = None,
    current_user: Optional[Dict[str, Any]] = Depends(get_current_user),
    session: AsyncSession = Depends(get_db)
):
    """
    Отмена заказа пользователем. Заказ можно отменить только если его статус позволяет отмену.
    
    - **order_id**: ID заказа
    - **cancel_reason**: Причина отмены
    """
    try:
        # Проверяем, авторизован ли пользователь
        if not current_user:
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Для отмены заказа необходима авторизация",
                headers={"WWW-Authenticate": "Bearer"},
            )
            
        # Получаем ID пользователя из токена
        user_id = current_user["user_id"]
        is_admin = current_user.get("is_admin", False)
        
        # Отменяем заказ
        order = await cancel_order(
            session=session,
            order_id=order_id,
            user_id=user_id,
            is_admin=is_admin,
            cancel_reason=cancel_reason
        )
        
        if not order:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Заказ с ID {order_id} не найден",
            )
        
        # Коммитим сессию
        await session.commit()
        
        # Вручную запрашиваем заказ со всеми связанными данными
        loaded_order = await get_order_by_id(session, order.id)
        
        # Инвалидируем кэш заказа
        await invalidate_order_cache(order_id)
        # Инвалидируем кэш статистики
        await invalidate_statistics_cache()
        logger.info("Кэш заказа %s и статистики инвалидирован после отмены заказа", order_id)
        
        # Преобразуем модель в схему
        return OrderResponse.model_validate(loaded_order)
    except ValueError as e:
        logger.error("Ошибка при отмене заказа: %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(e),
        ) from e
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Непредвиденная ошибка при отмене заказа: %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Произошла ошибка при отмене заказа",
        ) from e

@router.post("/{order_id}/reorder", status_code=status.HTTP_201_CREATED)
async def reorder_endpoint(
    order_id: int = Path(..., ge=1),
    personal_data_agreement: bool = Body(..., embed=True, description="Согласие на обработку персональных данных"),
    current_user: Dict[str, Any] = Depends(get_current_user),
    session: AsyncSession = Depends(get_db)
):
    """
    Повторение заказа - создание нового заказа с теми же товарами.
    
    - **order_id**: ID заказа для повторения
    """
    try:
        # Проверяем авторизацию пользователя
        if not current_user or not current_user.get("user_id"):
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Для повторения заказа необходимо авторизоваться"
            )
        # Проверяем согласие на обработку персональных данных
        if not personal_data_agreement:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Для повторения заказа необходимо согласие на обработку персональных данных"
            )
        
        user_id = current_user.get("user_id")
        token = current_user.get("token")
        
        # Получаем исходный заказ по ID
        original_order = await get_order_by_id(session, order_id)
        if not original_order:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Заказ с ID {order_id} не найден"
            )
        
        # Проверяем, принадлежит ли заказ текущему пользователю
        if original_order.user_id != user_id:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="У вас нет прав для повторения этого заказа"
            )
        
        # Получаем список товаров из заказа
        items = []
        product_ids = []
        for item in original_order.items:
            items.append({
                "product_id": item.product_id,
                "quantity": item.quantity
            })
            product_ids.append(item.product_id)
        
        # Проверяем доступность товаров
        availability = await check_products_availability(product_ids)
        
        # Проверяем, все ли товары доступны
        unavailable_products = [pid for pid, available in availability.items() if not available]
        if unavailable_products:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Товары с ID {unavailable_products} больше недоступны для заказа"
            )
        
        # Получаем информацию о товарах
        products_info = await get_products_info(product_ids, token)
        
        # Создаем новые данные для заказа на основе старого
        from schemas import OrderCreate, OrderItemCreate
        
        # Преобразуем товары в формат для создания заказа
        order_items = []
        for item in original_order.items:
            order_items.append(OrderItemCreate(
                product_id=item.product_id,
                quantity=item.quantity
            ))
        
        new_order_data = OrderCreate(
            full_name=original_order.full_name,
            email=original_order.email,
            phone=original_order.phone,
            region=original_order.region,
            city=original_order.city,
            street=original_order.street,
            comment=f"Повторный заказ на основе заказа #{original_order.id}",
            personal_data_agreement=personal_data_agreement,
            items=order_items
        )
        
        # Создаем новый заказ
        new_order = await create_order(
            session=session,
            user_id=user_id,
            order_data=new_order_data,
            product_service_url=PRODUCT_SERVICE_URL,
            token=token
        )
        
        # Коммитим сессию
        await session.commit()
        
        # Получаем созданный заказ
        created_order = await get_order_by_id(session, new_order.id)
        
        # Возвращаем информацию о созданном заказе
        response_data = OrderResponse.model_validate(created_order)
        return {
            "success": True,
            "message": "Заказ успешно повторен",
            "order": response_data,
            "order_id": new_order.id
        }
    except HTTPException:
        raise
    except ValueError as e:
        logger.error("Ошибка при повторении заказа: %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(e)
        ) from e
    except Exception as e:
        logger.error("Непредвиденная ошибка при повторении заказа: %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Произошла ошибка при повторении заказа"
        ) from e

# Административные маршруты
@admin_router.get("", response_model=PaginatedResponse,dependencies=[Depends(get_admin_user)])
async def list_all_orders(
    filters = Depends(get_order_filter_params),
    session: AsyncSession = Depends(get_db)
):
    """
    Получение списка всех заказов (только для администраторов).
    
    Поддерживает фильтрацию по различным параметрам.
    """
    try:
        # Создаем ключ для кэша на основе параметров фильтрации
        filter_params = filters.model_dump_json()
        cache_key = hashlib.md5(filter_params.encode()).hexdigest()
        
        # Пытаемся получить данные из кэша
        cached_orders = await get_cached_orders_list(cache_key)
        if cached_orders:
            logger.info("Данные о всех заказах получены из кэша по ключу %s", cache_key)
            return cached_orders
        
        # Если данных нет в кэше, получаем из БД
        orders, total = await get_orders(
            session=session,
            filters=filters
        )
        
        # Формируем ответ с преобразованием моделей в схемы
        order_responses = []
        for order in orders:
            # Создаем базовый объект OrderResponse без промокода
            order_response = OrderResponse.model_validate(order)
            
            # Если у заказа есть промокод, создаем OrderResponseWithPromo
            if order.promo_code_id:
                # Создаем расширенный ответ с информацией о промокоде
                with_promo = OrderResponseWithPromo(**order_response.model_dump())
                
                try:
                    # Получаем промокод и устанавливаем его
                    promo_code = await session.get(PromoCodeModel, order.promo_code_id)
                    if promo_code:
                        with_promo.promo_code = PromoCodeResponse.model_validate(promo_code)
                        logger.info("Для заказа %s загружен промокод %s (в админ списке)", order.id, promo_code.code)
                except (ValueError, AttributeError, KeyError) as e:
                    logger.warning("Не удалось загрузить промокод %s для заказа %s: %s", order.promo_code_id, order.id, str(e))
                
                order_responses.append(with_promo)
            else:
                # Заказ без промокода, используем обычный OrderResponse
                order_responses.append(order_response)
        
        response = PaginatedResponse(
            items=order_responses,
            total=total,
            page=filters.page,
            size=filters.size,
            pages=0  # Будет вычислено в валидаторе
        )
        
        # Кэшируем результат
        await cache_orders_list(cache_key, response)
        logger.info("Данные о всех заказах добавлены в кэш по ключу %s", cache_key)
        
        return response
    except Exception as e:
        logger.error("Ошибка при получении списка заказов: %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Произошла ошибка при получении списка заказов",
        ) from e

@admin_router.get("/statistics", response_model=OrderStatistics,dependencies=[Depends(get_admin_user)])
async def get_admin_orders_statistics(
    session: AsyncSession = Depends(get_db),
):
    """
    Получение статистики по всем заказам (только для администраторов)
    """
    try:
        # Пытаемся получить данные из кэша
        cached_statistics = await get_cached_order_statistics()
        if cached_statistics:
            logger.info("Статистика всех заказов получена из кэша")
            return cached_statistics
        
        # Если данных нет в кэше, получаем из БД
        statistics = await get_order_statistics(session)
        
        # Кэшируем результат
        await cache_order_statistics(statistics)
        logger.info("Статистика всех заказов добавлена в кэш")
        
        return statistics
    except Exception as e:
        logger.error("Ошибка при получении статистики заказов: %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Ошибка при получении статистики заказов"
        ) from e

@admin_router.get("/statistics/report", response_model=OrderStatistics,dependencies=[Depends(get_admin_user)])
async def get_admin_orders_statistics_by_date(
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
    session: AsyncSession = Depends(get_db),
):
    """
    Получение статистики по заказам за указанный период (только для администраторов)
    
    - **date_from**: Дата начала периода в формате YYYY-MM-DD
    - **date_to**: Дата окончания периода в формате YYYY-MM-DD
    """
    try:
        # Формируем ключ кэша на основе параметров запроса
        cache_key = f"order_statistics:report:{date_from or 'all'}:{date_to or 'all'}"
        
        # Пытаемся получить данные из кэша
        cached_statistics = await get_cached_data(cache_key)
        if cached_statistics:
            logger.info("Статистика заказов за период получена из кэша")
            return cached_statistics
        
        # Если данных нет в кэше, получаем из БД
        statistics = await get_order_statistics_by_date(session, date_from, date_to)
        
        # Кэшируем результат на короткое время (5 минут)
        await set_cached_data(cache_key, statistics, 300)
        logger.info("Статистика заказов за период добавлена в кэш")
        
        return statistics
    except Exception as e:
        logger.error("Ошибка при получении статистики заказов за период: %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Ошибка при получении статистики заказов за период"
        ) from e

@admin_router.get("/{order_id}", response_model=OrderDetailResponseWithPromo,dependencies=[Depends(get_admin_user)])
async def get_order_admin(
    order_id: int = Path(..., ge=1),
    session: AsyncSession = Depends(get_db)
):
    """
    Получение информации о заказе по ID (только для администраторов)
    
    - **order_id**: ID заказа
    """
    try:
        # Пытаемся получить данные из кэша
        cached_order = await get_cached_order(order_id, admin=True)
        if cached_order:
            logger.info("Данные о заказе %s получены из кэша (админ)", order_id)
            return cached_order
            
        # Получаем заказ
        order = await get_order_by_id(session, order_id)
        
        if not order:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Заказ с ID {order_id} не найден",
            )
        
        # Преобразуем модель в схему без промокода
        order_response = OrderDetailResponse.model_validate(order)
        
        # Создаем расширенный ответ с возможностью добавления промокода
        response_with_promo = OrderDetailResponseWithPromo(**order_response.model_dump())
        
        # Вручную обрабатываем промокод
        if order.promo_code_id:
            try:
                promo_code = await session.get(PromoCodeModel, order.promo_code_id)
                if promo_code:
                    response_with_promo.promo_code = PromoCodeResponse.model_validate(promo_code)
                    logger.info("Для заказа %s загружен промокод %s (админ)", order.id, promo_code.code)
            except (ValueError, AttributeError, KeyError) as e:
                logger.warning("Не удалось загрузить промокод %s для заказа %s: %s", order.promo_code_id, order.id, str(e))
        
        # Кэшируем результат с админским ключом
        cache_key = f"{CacheKeys.ORDER_PREFIX}{order_id}:admin"
        await cache_order(order_id, response_with_promo, admin=True, cache_key=cache_key)
        logger.info("Данные о заказе %s добавлены в кэш (админ)", order_id)
        
        return response_with_promo
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Ошибка при получении информации о заказе (админ): %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Произошла ошибка при получении информации о заказе",
        ) from e

@admin_router.put("/{order_id}", response_model=OrderResponse)
async def update_order_admin(
    order_id: int = Path(..., ge=1),
    order_data: OrderUpdate = Body(...),
    current_user: Dict[str, Any] = Depends(get_admin_user),
    session: AsyncSession = Depends(get_db)
):
    """
    Обновление информации о заказе (только для администраторов).
    
    - **order_id**: ID заказа
    - **order_data**: Данные для обновления
    """
    try:
        # Обновляем заказ
        order = await update_order(
            session=session,
            order_id=order_id,
            order_data=order_data
        )
        
        if not order:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Заказ с ID {order_id} не найден",
            )
        
        # Добавляем запись в историю статусов, если статус был изменен
        if order_data.status_id is not None:
            # Используем комментарий, если он предоставлен, иначе используем стандартное сообщение
            status_note = order_data.comment if order_data.comment else "Статус обновлен администратором"
            
            await OrderStatusHistoryModel.add_status_change(
                session=session,
                order_id=order.id,
                status_id=order_data.status_id,
                changed_by_user_id=current_user["user_id"],
                notes=status_note
            )
        
        # Коммитим сессию
        await session.commit()
        
        # Вручную запрашиваем заказ со всеми связанными данными
        loaded_order = await get_order_by_id(session, order.id)
        
        # Инвалидируем кэш заказа
        await invalidate_order_cache(order_id)
        # Инвалидируем кэш статистики, если изменился статус
        if order_data.status_id is not None:
            await invalidate_statistics_cache()
        logger.info("Кэш заказа %s инвалидирован после обновления заказа", order_id)
        
        # Преобразуем модель в схему
        return OrderResponse.model_validate(loaded_order)
    except ValueError as e:
        logger.error("Ошибка при обновлении заказа: %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(e),
        ) from e
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Непредвиденная ошибка при обновлении заказа: %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Произошла ошибка при обновлении заказа",
        ) from e

@admin_router.post("/{order_id}/status", response_model=OrderResponse)
async def update_order_status(
    order_id: int = Path(..., ge=1),
    status_data: dict = Body(...),
    current_user: Dict[str, Any] = Depends(get_admin_user),
    session: AsyncSession = Depends(get_db)
):
    """
    Обновление статуса заказа (только для администраторов).
    Упрощенный эндпоинт для обновления только статуса.
    
    - **order_id**: ID заказа
    - **status_data**: Данные о новом статусе:
      - status_id: ID нового статуса
      - comment: Комментарий к изменению статуса (опционально)
    """
    try:
        logger.info("Запрос на обновление статуса заказа. ID: %s, данные: %s", order_id, status_data)
        
        # Проверяем наличие обязательного поля status_id
        if not status_data.get("status_id"):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Не указан ID статуса",
            )
        
        # Создаем объект OrderUpdate только с нужными полями
        status_id = status_data.get("status_id")
        comment = status_data.get("comment")
        
        # Проверяем, что статус существует
        order_status = await session.get(OrderStatusModel, status_id)
        if not order_status:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Статус с ID {status_id} не найден",
            )
        
        # Получаем текущий заказ
        order = await get_order_by_id(session, order_id)
        if not order:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Заказ с ID {order_id} не найден",
            )
        
        # Обновляем статус заказа
        order.status_id = status_id
        session.add(order)
        
        # Добавляем запись в историю статусов
        status_note = comment if comment else "Статус обновлен администратором"
        await OrderStatusHistoryModel.add_status_change(
            session=session,
            order_id=order_id,
            status_id=status_id,
            changed_by_user_id=current_user["user_id"],
            notes=status_note
        )
        
        # Получаем информацию о новом статусе
        new_status = await session.get(OrderStatusModel, status_id)
        new_status_name = new_status.name if new_status else "Неизвестный"
        
        # Получаем информацию о старом статусе
        old_status = await session.get(OrderStatusModel, order.status_id)
        old_status_name = old_status.name if old_status else "Неизвестный"
        
        # Фиксируем изменения в базе данных
        await session.commit()
        
        # Инвалидируем кэш заказа
        await invalidate_order_cache(order_id)
        # Инвалидируем кэш статистики
        await invalidate_statistics_cache()
        logger.info("Кэш заказа %s и статистики инвалидирован после изменения статуса с '%s' на '%s'", order_id, old_status_name, new_status_name)
        
        # Обновляем заказ в сессии
        updated_order = await get_order_by_id(session, order_id)

        # Если у заказа есть промокод, загружаем его ПЕРЕД отправкой уведомления
        if updated_order.promo_code_id:
            # Загружаем промокод
            promo_code = await session.get(PromoCodeModel, updated_order.promo_code_id)
            if promo_code:
                # Создаем словарь с данными промокода для передачи в RabbitMQ
                updated_order.promo_code_dict = {
                    "code": promo_code.code,
                    "discount_percent": promo_code.discount_percent or 0
                }
                logger.info("Для заказа %s загружен промокод %s (при обновлении статуса)", updated_order.id, promo_code.code)
        
        # Отправляем уведомление об изменении статуса через Notifications Service
        try:
            # Проверяем, принадлежит ли заказ авторизованному пользователю
            if updated_order.user_id:
                # Для заказов авторизованных пользователей всегда проверяем настройки уведомлений
                logger.info("Отправка уведомления об изменении статуса для авторизованного пользователя: %s", updated_order.user_id)
                await check_notification_settings(updated_order.user_id, "order.status_changed", updated_order.id)
            else:
                # Для заказов неавторизованных пользователей проверяем флаг согласия на уведомления
                if updated_order.receive_notifications and updated_order.email:
                    logger.info("Отправка уведомления об изменении статуса для неавторизованного пользователя на email: %s", updated_order.email)
                    # Используем специальный метод для неавторизованных пользователей
                    await check_notification_settings(None, "order.status_changed", updated_order.id)
                else:
                    logger.info("Уведомление об изменении статуса не отправлено: неавторизованный пользователь не дал согласие или не указал email")
            
            logger.info("Отправлено событие 'order.status_changed' для заказа %s", order_id)
        except (ConnectionError, TimeoutError, ValueError) as e:
            logger.error("Ошибка при отправке события изменения статуса: %s", e)
        
        return OrderResponse.model_validate(updated_order)
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Ошибка при обновлении статуса заказа: %s", str(e))
        await session.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Произошла ошибка при обновлении статуса заказа: {str(e)}",
        ) from e

@admin_router.post("/{order_id}/items", response_model=OrderItemsUpdateResponse)
async def update_order_items_endpoint(
    order_id: int,
    items_data: OrderItemsUpdate,
    session: AsyncSession = Depends(get_db),
    current_user: Dict[str, Any] = Depends(get_admin_user)
):
    """Обновление товаров в заказе (админский доступ)"""
    logger.info("Запрос на обновление элементов заказа %s: items_to_add=%s items_to_update=%s items_to_remove=%s", order_id, items_data.items_to_add, items_data.items_to_update, items_data.items_to_remove)
    
    try:
        # Получаем текущий статус заказа
        order = await get_order_by_id(session, order_id)
        if not order:
            raise HTTPException(status_code=404, detail="Заказ не найден")
        
        # Проверка статуса заказа - запрещаем редактировать товары для определенных статусов
        non_editable_statuses = ["Отправлен", "Доставлен", "Отменен", "Оплачен"]
        if order.status and order.status.name in non_editable_statuses:
            raise HTTPException(
                status_code=400, 
                detail=f"Редактирование товаров невозможно для заказа в статусе '{order.status.name}'"
            )
        
        # Подготавливаем данные для обновления
        items_to_add = [item.model_dump() for item in items_data.items_to_add] if items_data.items_to_add else []
        items_to_update = items_data.items_to_update if items_data.items_to_update else {}
        items_to_remove = items_data.items_to_remove if items_data.items_to_remove else []
        
        # Вызываем сервисную функцию
        updated_order = await update_order_items(
            order_id=order_id,
            items_to_add=items_to_add,
            items_to_update=items_to_update,
            items_to_remove=items_to_remove,
            session=session,
            user_id=int(current_user.get("sub")) if current_user.get("sub") else None
        )
        
        if not updated_order:
            raise HTTPException(status_code=400, detail="Ошибка при обновлении элементов заказа")
        
        # Создаем ответ с данными заказа
        order_dict = {
            "id": updated_order.id,
            "user_id": updated_order.user_id,
            "status_id": updated_order.status_id,
            "status": {
                "id": updated_order.status.id,
                "name": updated_order.status.name,
                "description": updated_order.status.description,
                "color": updated_order.status.color,
                "allow_cancel": updated_order.status.allow_cancel,
                "is_final": updated_order.status.is_final,
                "sort_order": updated_order.status.sort_order
            },
            "created_at": updated_order.created_at,
            "updated_at": datetime.now(),
            "total_price": updated_order.total_price,
            "promo_code_id": updated_order.promo_code_id,
            "discount_amount": updated_order.discount_amount,
            "full_name": updated_order.full_name,
            "email": updated_order.email,
            "phone": updated_order.phone,
            "region": updated_order.region,
            "city": updated_order.city,
            "street": updated_order.street,
            "comment": updated_order.comment,
            "is_paid": updated_order.is_paid,
            "items": [
                {
                    "id": item.id,
                    "order_id": item.order_id,
                    "product_id": item.product_id,
                    "product_name": item.product_name,
                    "product_price": item.product_price,
                    "quantity": item.quantity,
                    "total_price": item.total_price
                } 
                for item in updated_order.items
            ],
            "order_number": f"{updated_order.id}-{updated_order.created_at.year}"
        }
        
        # Добавляем данные о промокоде, если он есть
        if hasattr(updated_order, 'promo_code') and updated_order.promo_code:
            promo = updated_order.promo_code
            order_dict["promo_code"] = {
                "id": promo.id,
                "code": promo.code,
                "discount_percent": promo.discount_percent,
                "discount_amount": promo.discount_amount,
                "valid_until": promo.valid_until,
                "is_active": promo.is_active,
                "created_at": promo.created_at,
                "updated_at": datetime.now()
            }
        
        return OrderItemsUpdateResponse(success=True, order=order_dict)
    except (HTTPException, ValueError, AttributeError, KeyError) as e:
        logger.error("Ошибка при обновлении элементов заказа: %s", str(e))
        return OrderItemsUpdateResponse(
            success=False, 
            errors={"message": f"Ошибка при обновлении элементов заказа: {str(e)}"}
        )

@admin_router.post("/batch-status", response_model=List[OrderResponse])
async def update_batch_status(
    update_data: BatchStatusUpdate,
    current_user: Dict[str, Any] = Depends(get_admin_user),
    session: AsyncSession = Depends(get_db)
):
    """
    Обновление статуса для нескольких заказов одновременно (только для администраторов).
    
    - **update_data**: Данные для массового обновления статусов
    """
    try:
        # Проверяем существование статуса
        order_status = await session.get(OrderStatusModel, update_data.status_id)
        if not order_status:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Статус с ID {update_data.status_id} не найден",
            )
        
        updated_orders = []
        for order_id in update_data.order_ids:
            try:
                # Создаем данные для обновления
                order_data = OrderUpdate(status_id=update_data.status_id)
                
                # Обновляем заказ
                order = await update_order(
                    session=session,
                    order_id=order_id,
                    order_data=order_data
                )
                
                if order:
                    # Добавляем запись в историю статусов
                    await OrderStatusHistoryModel.add_status_change(
                        session=session,
                        order_id=order.id,
                        status_id=update_data.status_id,
                        changed_by_user_id=current_user["user_id"],
                        notes=update_data.notes or "Массовое обновление статуса"
                    )
                    
                    loaded_order = await get_order_by_id(session, order.id)
                    updated_orders.append(loaded_order)
            except (ValueError, HTTPException, AttributeError) as e:
                logger.error("Ошибка при обновлении заказа %s: %s", order_id, str(e))
                # Продолжаем с другими заказами
        
        # Коммитим сессию
        await session.commit()
        
        # Преобразуем модели в схемы
        return [OrderResponse.model_validate(order) for order in updated_orders]
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Непредвиденная ошибка при массовом обновлении статусов: %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Произошла ошибка при массовом обновлении статусов",
        ) from e

@router.post("/check-can-review", status_code=status.HTTP_200_OK)
async def check_can_review(
    request_data: Dict[str, Any] = Body(...),
    session: AsyncSession = Depends(get_db)
):
    """
    Проверка, может ли пользователь оставить отзыв на товар
    (заказал товар и заказ доставлен)
    """
    try:
        user_id = request_data.get("user_id")
        product_id = request_data.get("product_id")
        
        if not user_id or not product_id:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Необходимо указать user_id и product_id"
            )
        
        # Проверяем, есть ли завершенные заказы с этим товаром у пользователя
        query = text("""
            SELECT COUNT(*) FROM orders o
            JOIN order_items oi ON o.id = oi.order_id
            WHERE o.user_id = :user_id 
            AND oi.product_id = :product_id
            AND o.status_id = 5
        """)
        
        result = await session.execute(
            query, 
            {"user_id": user_id, "product_id": product_id}
        )
        count = result.scalar_one()
        
        return {"can_review": count > 0}
    except Exception as e:
        logger.error("Ошибка при проверке возможности оставить отзыв: %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Ошибка при проверке возможности оставить отзыв"
        ) from e

@router.post("/check-can-review-store", status_code=status.HTTP_200_OK)
async def check_can_review_store(
    request_data: Dict[str, Any] = Body(...),
    session: AsyncSession = Depends(get_db)
):
    """
    Проверка, может ли пользователь оставить отзыв на магазин
    (имеет хотя бы один завершенный заказ)
    """
    try:
        user_id = request_data.get("user_id")
        
        if not user_id:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Необходимо указать user_id"
            )
        
        # Проверяем, есть ли завершенные заказы у пользователя
        query = text("""
            SELECT COUNT(*) FROM orders o
            WHERE o.user_id = :user_id 
            AND o.status_id = 5
        """)
        
        result = await session.execute(query, {"user_id": user_id})
        count = result.scalar_one()
        
        return {"can_review": count > 0}
    except Exception as e:
        logger.error("Ошибка при проверке возможности оставить отзыв на магазин: %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Ошибка при проверке возможности оставить отзыв на магазин"
        ) from e


@admin_router.post("", response_model=OrderResponseWithPromo, status_code=status.HTTP_201_CREATED)
async def create_order_admin(
    order_data: AdminOrderCreate,
    current_user: Dict[str, Any] = Depends(get_admin_user),
    session: AsyncSession = Depends(get_db)
):
    """
    Создание нового заказа из админки. Доступно только для администраторов.
    
    - **order_data**: Данные для создания заказа
    """
    try:
        logger.info("Получен запрос на создание заказа из админки: %s", order_data)
        
        # Получаем токен авторизации администратора для запросов к другим сервисам
        token = current_user.get("token")
        
        # Если передан user_id, используем его, иначе заказ создается без привязки к пользователю
        user_id = order_data.user_id

        # Обработка пустого промокода
        promo_code = order_data.promo_code
        if promo_code == "" or (promo_code and len(promo_code) < 3):
            promo_code = None
            
        # Нормализация телефона
        phone = order_data.phone
        # Добавляем префикс, если его нет
        if phone and not (phone.startswith('8') or phone.startswith('+7')):
            phone = '8' + phone
        
        # Убеждаемся, что телефон имеет нужную длину
        if phone and len(phone) < 11:
            if phone.startswith('8'):
                # Дополняем до 11 цифр для формата 8XXXXXXXXXX
                missing_digits = 11 - len(phone)
                if missing_digits > 0:
                    phone = phone + '0' * missing_digits
            elif phone.startswith('+7'):
                # Дополняем до 12 цифр для формата +7XXXXXXXXXX
                missing_digits = 12 - len(phone)
                if missing_digits > 0:
                    phone = phone + '0' * missing_digits

        # Преобразуем данные из AdminOrderCreate в OrderCreate
        create_data = OrderCreate(
            items=order_data.items,
            full_name=order_data.full_name,
            email=order_data.email,
            phone=phone,  # Используем нормализованный телефон
            region=order_data.region,
            city=order_data.city,
            street=order_data.street,
            comment=order_data.comment,
            promo_code=promo_code,  # Используем обработанный промокод
            personal_data_agreement=True  # Для админа всегда True
        )
        
        # Проверяем наличие товаров
        product_ids = [item.product_id for item in order_data.items]
        availability = await check_products_availability(product_ids)
        
        # Проверяем, все ли товары доступны
        unavailable_products = [pid for pid, available in availability.items() if not available]
        if unavailable_products:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Товары с ID {unavailable_products} недоступны для заказа",
            )
        
        # Создаем заказ через общую функцию
        order = await create_order(
            session=session,
            user_id=user_id,
            order_data=create_data,
            product_service_url=PRODUCT_SERVICE_URL,
            token=token
        )
        
        # Если указан начальный статус и он отличается от дефолтного
        if order_data.status_id:
            try:
                status_data = OrderStatusHistoryCreate(
                    status_id=order_data.status_id,
                    notes="Статус установлен при создании заказа администратором"
                )
                order = await change_order_status(
                    session=session,
                    order_id=order.id,
                    status_data=status_data,
                    user_id=current_user["user_id"],
                    is_admin=True
                )
                logger.info("Для нового заказа %s установлен статус %s", order.id, order_data.status_id)
            except (ValueError, HTTPException, AttributeError) as e:
                logger.warning("Не удалось установить начальный статус %s для заказа %s: %s", order_data.status_id, order.id, str(e))
        
        # Устанавливаем флаг оплаты, если указан
        if order_data.is_paid:
            order.is_paid = True
            await session.commit()
            logger.info("Для нового заказа %s установлен флаг оплаты", order.id)
        
        # Явно коммитим сессию, чтобы убедиться, что все связанные данные загружены
        await session.commit()
        
        # Вручную запрашиваем заказ со всеми связанными данными
        loaded_order = await get_order_by_id(session, order.id)
        
        # Если у заказа есть промокод, загружаем его
        if loaded_order.promo_code_id:
            # Загружаем промокод
            promo_code = await session.get(PromoCodeModel, loaded_order.promo_code_id)
            if promo_code:
                # Создаем словарь с данными промокода для передачи в RabbitMQ
                loaded_order.promo_code_dict = {
                    "code": promo_code.code,
                    "discount_percent": promo_code.discount_percent or 0
                }
                logger.info("Для нового заказа %s загружен промокод %s", loaded_order.id, promo_code.code)
        
        # Отправляем подтверждение заказа на email через Notifications Service
        try:
            if order_data.email and user_id != None:
                logger.info("Отправка подтверждения заказа на email: %s", order_data.email)
                await check_notification_settings(loaded_order.user_id, "order.created", loaded_order.id)
        except (ConnectionError, TimeoutError, ValueError) as e:
            logger.error("Ошибка при отправке email о заказе: %s", str(e))
        
        # Явно инвалидируем кэш заказов перед возвратом ответа
        await invalidate_order_cache(order.id)
        logger.info("Кэш заказа %s и связанных списков инвалидирован перед возвратом ответа", order.id)
        
        # Преобразуем модель в схему без промокода
        order_response = OrderResponse.model_validate(loaded_order)
        
        # Создаем расширенный ответ с возможностью добавления промокода
        response_with_promo = OrderResponseWithPromo(**order_response.model_dump())
        
        # Если у заказа есть промокод, загружаем его данные вручную
        if loaded_order.promo_code_id:
            try:
                promo_code = await session.get(PromoCodeModel, loaded_order.promo_code_id)
                if promo_code:
                    response_with_promo.promo_code = PromoCodeResponse.model_validate(promo_code)
                    logger.info("Для нового заказа %s загружен промокод %s", loaded_order.id, promo_code.code)
            except (ValueError, AttributeError, KeyError) as e:
                logger.warning("Не удалось загрузить промокод %s для заказа %s: %s", loaded_order.promo_code_id, loaded_order.id, str(e))
        
        return response_with_promo
    except ValueError as e:
        logger.error("Ошибка при создании заказа из админки: %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=str(e),
        ) from e
    except Exception as e:
        logger.error("Непредвиденная ошибка при создании заказа из админки: %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Произошла ошибка при создании заказа",
        ) from e


@router.get("/{order_id}/service", response_model=OrderDetailResponseWithPromo, dependencies=[Depends(verify_service_jwt)])
async def get_order_service(
    order_id: int = Path(..., ge=1),
    session: AsyncSession = Depends(get_db)
):
    """
    Получение информации о заказе для сервисных нужд.
    
    - **order_id**: ID заказа
    """
    try:
        # Для сервисного запроса используем простой ключ без user_id
        cache_key = f"{CacheKeys.ORDER_PREFIX}{order_id}:service"
        cached_order = await get_cached_data(cache_key)
        if cached_order:
            logger.info("Данные о заказе %s получены из кэша (сервисный запрос)", order_id)
            return cached_order
            
        # Получаем заказ без проверки на принадлежность пользователю
        order = await get_order_by_id(
            session=session,
            order_id=order_id,
            user_id=None  # Для сервисных запросов получаем заказ без проверки пользователя
        )
        
        if not order:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Заказ с ID {order_id} не найден",
            )
        
        # Преобразуем модель в схему без промокода
        order_response = OrderDetailResponse.model_validate(order)
        
        # Создаем расширенный ответ с возможностью добавления промокода
        response_with_promo = OrderDetailResponseWithPromo(**order_response.model_dump())
        
        # Получаем только последний статус, если история статусов существует
        if hasattr(response_with_promo, "status_history") and response_with_promo.status_history:
            # Сортируем историю статусов по дате изменения (вместо created_at используем changed_at)
            sorted_history = sorted(
                response_with_promo.status_history,
                key=lambda x: x.changed_at,
                reverse=True
            )
            response_with_promo.status_history = [sorted_history[0]] if sorted_history else []
        
        # Вручную обрабатываем промокод
        if order.promo_code_id:
            try:
                promo_code = await session.get(PromoCodeModel, order.promo_code_id)
                if promo_code:
                    response_with_promo.promo_code = PromoCodeResponse.model_validate(promo_code)
                    logger.info("Для заказа %s загружен промокод %s", order.id, promo_code.code)
            except (ValueError, AttributeError, KeyError) as e:
                logger.warning("Не удалось загрузить промокод %s для заказа %s: %s", order.promo_code_id, order.id, str(e))
        
        # Кэшируем результат со специальным ключом для сервисных запросов
        await set_cached_data(cache_key, response_with_promo, DEFAULT_CACHE_TTL)
        logger.info("Данные о заказе %s добавлены в кэш для сервисных запросов", order_id)
        
        return response_with_promo
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Ошибка при получении информации о заказе: %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Произошла ошибка при получении информации о заказе",
        ) from e

@router.post("/{order_id}/unsubscribe", status_code=status.HTTP_200_OK)
async def unsubscribe_notifications(
    order_id: int = Path(..., ge=1),
    email: str = Body(..., embed=True),
    session: AsyncSession = Depends(get_db)
):
    """
    Отменить подписку на уведомления по email для неавторизованного пользователя.
    
    - **order_id**: ID заказа
    - **email**: Email, указанный при создании заказа для подтверждения
    """
    try:
        # Получаем заказ
        order = await get_order_by_id(session, order_id)
        
        if not order:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Заказ с ID {order_id} не найден",
            )
        
        # Проверяем принадлежность к пользователю
        if order.user_id:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Данный метод доступен только для заказов, созданных без регистрации",
            )
        
        # Проверяем email
        if order.email.lower() != email.lower():
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Указанный email не совпадает с email, указанным при создании заказа",
            )
        
        # Обновляем флаг получения уведомлений
        order.receive_notifications = False
        session.add(order)
        await session.commit()
        
        # Инвалидируем кэш заказа
        await invalidate_order_cache(order_id)
        logger.info("Кэш заказа %s инвалидирован после отмены подписки на уведомления", order_id)
        
        return {
            "success": True,
            "message": "Вы успешно отписались от уведомлений о заказе"
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Ошибка при отмене подписки на уведомления: %s", str(e))
        await session.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Произошла ошибка при отмене подписки на уведомления",
        ) from e

@admin_router.get("/reports/download", dependencies=[Depends(get_admin_user)])
async def generate_orders_report(
    format: str = Query(..., description="Формат отчета: csv, excel, pdf, word"),
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
    session: AsyncSession = Depends(get_db)
):
    """
    Генерация отчета по заказам в различных форматах
    
    - **format**: Формат отчета (csv, excel, pdf, word)
    - **date_from**: Дата начала периода (YYYY-MM-DD)
    - **date_to**: Дата окончания периода (YYYY-MM-DD)
    """
    try:
        # Получаем статистику за указанный период
        statistics = await get_order_statistics_by_date(session, date_from, date_to)
        
        # Создаем датафрейм со статистикой
        stats_data = {
            "Метрика": ["Всего заказов", "Общая сумма заказов", "Средняя стоимость заказа"],
            "Значение": [
                statistics.total_orders,
                statistics.total_revenue,
                round(statistics.average_order_value, 2)
            ]
        }
        
        # Добавляем статистику по статусам
        for status_name, count in statistics.orders_by_status.items():
            stats_data["Метрика"].append(f"Заказов со статусом '{status_name}'")
            stats_data["Значение"].append(count)
        
        df_stats = pd.DataFrame(stats_data)
        
        # Получаем список заказов за указанный период
        # Создаем фильтры для запроса заказов
        filters = OrderFilterParams(
            page=1,
            size=100,  # Максимально допустимое значение
            date_from=date_from,
            date_to=date_to
        )
        
        # Получаем заказы
        orders, _ = await get_orders(session, filters)
        
        # Создаем датафрейм с заказами
        orders_data = []
        for order in orders:
            orders_data.append({
                "ID": order.id,
                "Номер заказа": order.order_number,
                "Дата создания": order.created_at.strftime("%Y-%m-%d %H:%M:%S"),
                "Статус": order.status.name,
                "Клиент": order.full_name,
                "Email": order.email,
                "Телефон": order.phone,
                "Адрес": f"{order.region}, {order.city}, {order.street}",
                "Сумма заказа": order.total_price,
                "Скидка": order.discount_amount or 0,
                "Оплачен": "Да" if order.is_paid else "Нет"
            })
        
        df_orders = pd.DataFrame(orders_data)
        
        # Период для имени файла
        period = ""
        if date_from and date_to:
            period = f"_{date_from}_{date_to}"
        elif date_from:
            period = f"_{date_from}"
        elif date_to:
            period = f"_{date_to}"
        
        # Обработка в зависимости от формата
        if format.lower() == "csv":
            # Создаем буфер для CSV файла
            output = io.StringIO()
            
            # Записываем заголовок
            output.write("Отчет по заказам\n")
            if date_from and date_to:
                output.write(f"Период: с {date_from} по {date_to}\n\n")
            elif date_from:
                output.write(f"Период: с {date_from}\n\n")
            elif date_to:
                output.write(f"Период: по {date_to}\n\n")
            else:
                output.write("Период: все время\n\n")
            
            # Записываем статистику
            output.write("Общая статистика:\n")
            df_stats.to_csv(output, index=False, sep=';')
            
            output.write("\n\nСписок заказов:\n")
            df_orders.to_csv(output, index=False, sep=';')
            
            # Возвращаем StreamingResponse
            output.seek(0)
            
            # Форматируем имя файла                
            filename = f"orders_report{period}.csv"
            
            return StreamingResponse(
                iter([output.getvalue()]),
                media_type="text/csv",
                headers={"Content-Disposition": f"attachment; filename={filename}"}
            )
        
        elif format.lower() == "excel":
            # Создаем буфер для Excel файла
            output = io.BytesIO()
            
            # Создаем Excel файл
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                # Добавляем лист со статистикой
                df_stats.to_excel(writer, sheet_name='Статистика', index=False)
                
                # Добавляем лист с заказами
                df_orders.to_excel(writer, sheet_name='Заказы', index=False)
            
            # Возвращаем StreamingResponse
            output.seek(0)
            
            # Форматируем имя файла
            filename = f"orders_report{period}.xlsx"
            
            return StreamingResponse(
                io.BytesIO(output.getvalue()),
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                headers={"Content-Disposition": f"attachment; filename={filename}"}
            )
        
        elif format.lower() == "pdf":
            try:
                # Импортируем необходимые библиотеки
                from reportlab.lib import colors
                from reportlab.lib.pagesizes import letter, landscape
                from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet
                
                buffer = io.BytesIO()
                
                # Создаем PDF документ
                doc = SimpleDocTemplate(buffer, pagesize=landscape(letter))
                elements = []
                
                # Стили
                styles = getSampleStyleSheet()
                title_style = styles['Heading1']
                subtitle_style = styles['Heading2']
                normal_style = styles['Normal']
                
                # Заголовок
                elements.append(Paragraph("Отчет по заказам", title_style))
                
                # Период отчета
                if date_from and date_to:
                    elements.append(Paragraph(f"Период: с {date_from} по {date_to}", subtitle_style))
                elif date_from:
                    elements.append(Paragraph(f"Период: с {date_from}", subtitle_style))
                elif date_to:
                    elements.append(Paragraph(f"Период: по {date_to}", subtitle_style))
                else:
                    elements.append(Paragraph("Период: все время", subtitle_style))
                
                elements.append(Spacer(1, 20))
                
                # Общая статистика
                elements.append(Paragraph("Общая статистика", subtitle_style))
                
                # Создаем таблицу со статистикой
                stats_data = [['Метрика', 'Значение']]
                stats_data.append(['Всего заказов', str(statistics.total_orders)])
                stats_data.append(['Общая сумма заказов', f"{statistics.total_revenue:.2f} руб."])
                stats_data.append(['Средняя стоимость заказа', f"{statistics.average_order_value:.2f} руб."])
                
                # Добавляем статистику по статусам
                for status_name, count in statistics.orders_by_status.items():
                    stats_data.append([f"Заказов со статусом '{status_name}'", str(count)])
                
                # Создаем таблицу и задаем стиль
                stats_table = Table(stats_data, colWidths=[300, 150])
                stats_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('ALIGN', (1, 1), (1, -1), 'RIGHT'),
                ]))
                
                elements.append(stats_table)
                elements.append(Spacer(1, 20))
                
                # Список заказов
                elements.append(Paragraph("Список заказов", subtitle_style))
                
                # Создаем таблицу со списком заказов
                if orders:
                    # Заголовки таблицы
                    orders_data = [['ID', 'Номер заказа', 'Дата создания', 'Статус', 'Клиент', 'Сумма', 'Оплачен']]
                    
                    # Данные заказов
                    for order in orders:
                        orders_data.append([
                            str(order.id),
                            order.order_number,
                            order.created_at.strftime("%Y-%m-%d %H:%M"),
                            order.status.name,
                            order.full_name,
                            f"{order.total_price:.2f} руб.",
                            "Да" if order.is_paid else "Нет"
                        ])
                    
                    # Создаем таблицу и задаем стиль
                    orders_table = Table(orders_data, colWidths=[40, 80, 80, 80, 150, 70, 50])
                    orders_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('ALIGN', (0, 1), (0, -1), 'RIGHT'),
                        ('ALIGN', (5, 1), (5, -1), 'RIGHT'),
                        ('ALIGN', (6, 1), (6, -1), 'CENTER'),
                    ]))
                    
                    elements.append(orders_table)
                else:
                    elements.append(Paragraph("Нет заказов за выбранный период", normal_style))
                
                # Строим PDF документ
                doc.build(elements)
                
                # Готовим ответ
                buffer.seek(0)
                
                # Имя файла
                filename = f"orders_report{period}.pdf"
                
                return StreamingResponse(
                    buffer,
                    media_type="application/pdf",
                    headers={"Content-Disposition": f"attachment; filename={filename}"}
                )
            except Exception as e:
                logger.error(f"Ошибка при создании PDF отчета: {str(e)}")
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"Ошибка при создании PDF отчета: {str(e)}"
                )
        
        elif format.lower() == "word":
            try:
                # Импортируем необходимые библиотеки
                from docx import Document
                from docx.shared import Pt, Cm, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                # Создаем документ Word
                doc = Document()
                
                # Заголовок отчета
                title = doc.add_heading('Отчет по заказам', level=1)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Период отчета
                if date_from and date_to:
                    doc.add_heading(f"Период: с {date_from} по {date_to}", level=2)
                elif date_from:
                    doc.add_heading(f"Период: с {date_from}", level=2)
                elif date_to:
                    doc.add_heading(f"Период: по {date_to}", level=2)
                else:
                    doc.add_heading("Период: все время", level=2)
                
                # Добавляем раздел общей статистики
                doc.add_heading('Общая статистика', level=2)
                
                # Таблица со статистикой
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                
                # Заголовки таблицы
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Метрика'
                hdr_cells[1].text = 'Значение'
                
                # Форматирование заголовков
                for cell in table.rows[0].cells:
                    cell.paragraphs[0].runs[0].bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Данные статистики
                row = table.add_row().cells
                row[0].text = 'Всего заказов'
                row[1].text = str(statistics.total_orders)
                
                row = table.add_row().cells
                row[0].text = 'Общая сумма заказов'
                row[1].text = f"{statistics.total_revenue:.2f} руб."
                
                row = table.add_row().cells
                row[0].text = 'Средняя стоимость заказа'
                row[1].text = f"{statistics.average_order_value:.2f} руб."
                
                # Статистика по статусам
                for status_name, count in statistics.orders_by_status.items():
                    row = table.add_row().cells
                    row[0].text = f"Заказов со статусом '{status_name}'"
                    row[1].text = str(count)
                
                # Добавляем раздел списка заказов
                doc.add_heading('Список заказов', level=2)
                
                # Таблица заказов
                if orders:
                    table = doc.add_table(rows=1, cols=7)
                    table.style = 'Table Grid'
                    
                    # Заголовки таблицы
                    hdr_cells = table.rows[0].cells
                    headers = ['ID', 'Номер заказа', 'Дата создания', 'Статус', 'Клиент', 'Сумма', 'Оплачен']
                    for i, header in enumerate(headers):
                        hdr_cells[i].text = header
                    
                    # Форматирование заголовков
                    for cell in table.rows[0].cells:
                        cell.paragraphs[0].runs[0].bold = True
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Данные заказов
                    for order in orders:
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(order.id)
                        row_cells[1].text = order.order_number
                        row_cells[2].text = order.created_at.strftime("%Y-%m-%d %H:%M")
                        row_cells[3].text = order.status.name
                        row_cells[4].text = order.full_name
                        row_cells[5].text = f"{order.total_price:.2f} руб."
                        row_cells[6].text = "Да" if order.is_paid else "Нет"
                else:
                    doc.add_paragraph("Нет заказов за выбранный период")
                
                # Сохраняем документ в буфер
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                # Имя файла
                filename = f"orders_report{period}.docx"
                
                return StreamingResponse(
                    buffer,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": f"attachment; filename={filename}"}
                )
            except Exception as e:
                logger.error(f"Ошибка при создании Word отчета: {str(e)}")
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"Ошибка при создании Word отчета: {str(e)}"
                )
        
        else:
            # Возвращаем ошибку, если формат не поддерживается
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Неподдерживаемый формат отчета: {format}. Поддерживаемые форматы: csv, excel, pdf, word"
            )
    
    except Exception as e:
        logger.error("Ошибка при генерации отчета: %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Произошла ошибка при генерации отчета"
        ) from e